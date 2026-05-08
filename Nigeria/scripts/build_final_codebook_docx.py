"""Build the final research-grade coding codebook (.docx) for the Nigeria
Manfluencer Project.

Output:
    Codebooks/LLM Codebook/Nigeria Manfluencer Project - Final Coding Codebook.docx
"""
from __future__ import annotations
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).resolve().parents[2]
OUT = ROOT / "Codebooks" / "LLM Codebook" / "Nigeria Manfluencer Project - Final Coding Codebook.docx"
OUT.parent.mkdir(parents=True, exist_ok=True)

# ─── helpers ────────────────────────────────────────────────────────────────

def add_heading(doc, text, level=1, color=None):
    h = doc.add_heading(text, level=level)
    if color is not None:
        for r in h.runs:
            r.font.color.rgb = RGBColor(*color)
    return h

def para(doc, text, bold=False, italic=False, size=11, space_after=6):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    p.paragraph_format.space_after = Pt(space_after)
    return p

def bullet(doc, text, level=0, bold_lead=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_lead:
        r1 = p.add_run(bold_lead); r1.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    if level > 0:
        p.paragraph_format.left_indent = Inches(0.25 + 0.25*level)
    return p

def kv_table(doc, rows, col_widths=None):
    t = doc.add_table(rows=0, cols=2)
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    for k, v in rows:
        row = t.add_row()
        c0, c1 = row.cells
        c0.text = ""; c1.text = ""
        r = c0.paragraphs[0].add_run(k); r.bold = True; r.font.size = Pt(10)
        r2 = c1.paragraphs[0].add_run(v); r2.font.size = Pt(10)
    if col_widths:
        for row in t.rows:
            for cell, w in zip(row.cells, col_widths):
                cell.width = w
    return t

# ─── theme content ──────────────────────────────────────────────────────────

THEMES = [
    {
        "code": "Authority and Submission",
        "bundle": "Pro-traditional / regressive",
        "definition": "Comments or content asserting hierarchy between men and women — male headship, female submission, obedience, control, command, surname/lineage logic, or claims that women must surrender independence to men.",
        "when": "The text endorses, debates, or challenges unequal authority between men and women. Hierarchy or control language is central.",
        "when_not": "Use a different theme when family/marriage is the only setting and no hierarchy is asserted, or when the focus is on provider status without dominance framing.",
        "rule": "Requires explicit hierarchy or control vocabulary (lead, head, submit, obey, command, surname, control, serve).",
        "why": "Operationalizes patriarchal gender norms. Distinguishes ordinary family talk from hierarchy-building discourse, which is the core regressive frame in Shola's content and a substantial fraction of his audience.",
        "evidence": "10 audience + 7 content rows in the v2 audit (after corrections, ~12 audience). Ranked highly stable on the strict rule.",
        "creators": "Shola (8 audience primary), Agba and Banky audience as smaller share.",
    },
    {
        "code": "Male Victimhood",
        "bundle": "Pro-traditional / regressive",
        "definition": "Men framed as disadvantaged, exploited, or losing — child-support fraud, alimony scams, false-accusation panic, claims that the legal/cultural system is rigged against men.",
        "when": "Men appear as victims of women, marriage, courts, society, or feminism. Focus is on male loss.",
        "when_not": "Do not use when the framing is about women as threats in general (use Gender Grievance) or when content is about specific abuse cases (use Gender-Based Violence and Consent).",
        "rule": "Distinct from Gender Grievance: grievance frames women as scammers / threats; victimhood frames men as the harmed party. Both can co-occur.",
        "why": "Content codebook (Q7) lists 'Men are disadvantaged / victims' as a standalone option. Without this theme, Wizarab's child-support-fraud framing collapses into Gender Grievance and the male-loss register is lost.",
        "evidence": "Newly added in v3 after gap analysis against the content codebook. Predicted to fire on Wizarab content (child support, courts) and audience replies extending these frames.",
        "creators": "Wizarab (primary anchor), Agba secondary, audience replies on Shola.",
    },
    {
        "code": "Gender Grievance",
        "bundle": "Pro-traditional / regressive",
        "definition": "Generalized distrust of women, feminists, or modern women; gender-war framing; claims that equality is a scam, that women are opportunists, or that gender relations are zero-sum.",
        "when": "The text generalizes negatively about women / feminists or frames gender relations as zero-sum conflict.",
        "when_not": "Specific criticism of one named person without broader generalization. Use Sexual Morality if the focus is on individual moral behavior.",
        "rule": "Requires generalization across women as a class, not particular case.",
        "why": "Captures the affective fuel of manosphere-adjacent discourse. Identifies how women are framed as threats, cheats, or beneficiaries of unfair systems — central to Shola, Wizarab, and parts of Agba.",
        "evidence": "16 audience primary in v2; 18 after corrections. Stable at 78% precision with strict rule.",
        "creators": "Shola (13 audience primary), Wizarab content, Agba.",
    },
    {
        "code": "Sexual Morality",
        "bundle": "Pro-traditional / regressive",
        "definition": "Body count, cheating, abortion, pornography, BBL/body policing, female desirability and respectability, sexual double standards.",
        "when": "Sexual conduct or women's bodies are used to judge moral or marital worth.",
        "when_not": "When rape, consent, or abuse is central — use Gender-Based Violence and Consent. Generic infidelity discussions inside marriage advice may belong to Marriage and Family if morality is incidental.",
        "rule": "Focus is the moral framing of sexual behavior. Hierarchy of women based on sexual conduct is in scope.",
        "why": "Captures respectability politics and sexual double standards — a major entry point into misogynistic discourse. Agba's cheating-tolerance posts and Wizarab's BBL/body-shaming posts both fall here.",
        "evidence": "32 audience primary, 28 content. Stable at 93% precision.",
        "creators": "Agba (32 audience primary, exclusive), Wizarab content, Shola minor.",
    },
    {
        "code": "Relationship Tactics",
        "bundle": "Pro-traditional / regressive",
        "definition": "Tactical dating advice — scarcity, pursuit, availability, options, rejection, picking partners, attraction strategy, masculine-frame instructions.",
        "when": "The row gives or debates strategies for attracting, rejecting, controlling, or negotiating with partners.",
        "when_not": "Ordinary marriage talk, vague relationship praise, or general partnership discussions. Tactics must be tactical, not aspirational.",
        "rule": "Practical advice scripts, not values discussions. 'Don't be too available' is in scope; 'Marriage takes mutual respect' is not.",
        "why": "Key mechanism in short-form manosphere content: broad gender ideology often smuggled in as practical dating advice. Distinguishes tactical content from broader marriage talk.",
        "evidence": "33 audience, 47 content in v2. Tightened in v3 to 65% precision; expected to drop ~30% with strict rule.",
        "creators": "Shola (heavy content), Wizarab, Banky audience around dating discussions.",
    },
    {
        "code": "Provider and Status",
        "bundle": "Pro-traditional / regressive",
        "definition": "Money, income, career, status, respectability, masculine worth proven through provision or success.",
        "when": "The core claim links money / provision / income to whether a man is valued or whether a woman should respect or submit to him.",
        "when_not": "Money is incidental and the focus is submission, sexual morality, or partnership.",
        "rule": "Pressure to provide or economic value as masculine worth is the central frame.",
        "why": "Bridges self-help masculinity and gender hierarchy. Economic pressure becomes a reason to rank men, discipline women, or justify unequal relationships. Strong overlap with content codebook Q7 'provide/succeed' and Q8 'economic pressure'.",
        "evidence": "8 audience, 32 content. Most stable theme at 97.5% precision.",
        "creators": "Shola (16 content primary), Banky audience and content, Wizarab.",
    },
    {
        "code": "Male Accountability",
        "bundle": "Progressive",
        "definition": "Men must change, hold men accountable, men's own behavior as the problem to solve; refusal of deflection and counter-claims.",
        "when": "The text argues that men should be responsible for the harms men cause, that male behavior is the issue, or that men should hold each other accountable.",
        "when_not": "When sexual violence is the specific focus — use Gender-Based Violence and Consent (which subsumes consent advocacy). When the frame is about partnership reciprocity broadly — use Egalitarian Partnership.",
        "rule": "Distinct from Gender-Based Violence and Consent (which is about violence specifically) and from Egalitarian Partnership (which is about reciprocity in relationships).",
        "why": "Scope memo names this for Deyemi Okanlawon explicitly. Content codebook Q8 lists 'Men's behavior' as a problem-axis option. Without this theme, the men-must-change frame is lost or absorbed into adjacent codes.",
        "evidence": "Newly carved in v3 from prior bundling. Maps onto a non-trivial subset of Deyemi audience replies and Banky podcast segments.",
        "creators": "Deyemi (anchor), Banky audience.",
    },
    {
        "code": "Egalitarian Partnership",
        "bundle": "Progressive",
        "definition": "Mutual respect, shared money, shared parenting, listening, allyship, healthy reciprocity, non-dominating partnership.",
        "when": "The text explicitly promotes reciprocity, shared responsibility, or non-dominating partnership.",
        "when_not": "Generic praise of a relationship without naming a partnership norm. Specific advocacy for one gender (use Male Accountability or Gender-Based Violence and Consent).",
        "rule": "Requires a relational norm (shared, mutual, equal, listen) — not just praise of an individual.",
        "why": "Strongest positive counter-theme in the dataset and most useful for playbook / counternarrative development. Maps to content codebook Q9 'equality for women' solution and Q7 'equal partners' option.",
        "evidence": "9 audience, 10 content. Stable at 94.7% precision.",
        "creators": "Banky (anchor), Deyemi audience, Ebuka.",
    },
    {
        "code": "Gender-Based Violence and Consent",
        "bundle": "Progressive",
        "definition": "Rape, consent, abuse, victim stigma, false accusations, child abuse, molestation, prosecution, victim protection, legal/cultural accountability around violence.",
        "when": "Violence, consent, rape, abuse, or victim/accuser framing is the main issue.",
        "when_not": "Generic cheating or sexual shame without violence or consent. False-accusation rhetoric outside the violence-discourse context goes elsewhere.",
        "rule": "Scope is specifically violence and consent. Anti-rape advocacy and false-accusation pushback both code here as long as the discourse is rooted in the violence frame.",
        "why": "Memo names 'rape culture, victim protection' for Deyemi explicitly. The violence axis is analytically distinct from sexual morality (the moral-judgment axis).",
        "evidence": "59 audience primary in v2; 69 after corrections. Stable at 96.8% precision.",
        "creators": "Deyemi (audience anchor — 67 of 69 primary), Banky and Wizarab content.",
    },
    {
        "code": "Trauma and Mental Health",
        "bundle": "Progressive",
        "definition": "Male emotional expression, trauma, depression, healing, illness, grief, vulnerability, psychological harm, help-seeking.",
        "when": "Inner life, healing, emotional openness, or psychological harm is central.",
        "when_not": "The row says 'painful' rhetorically without engaging emotional content. Generic complaints about hardship without psychological framing.",
        "rule": "Emotional / mental life must be the substantive content, not figurative language.",
        "why": "Identifies healthier masculinity counternarratives around men's pain and help-seeking. Banky's testimony content and Deyemi's male-trauma frame both fall here. Content codebook Q7 'be emotionally open' and Q8 'mental health' map directly.",
        "evidence": "13 audience, 8 content. Stable at 71% precision (some figurative-language false positives).",
        "creators": "Banky (anchor), Deyemi audience and content.",
    },
    {
        "code": "Self-Discipline",
        "bundle": "Progressive",
        "definition": "Personal responsibility, maturity, restraint, growth, learning/unlearning, habit formation, decision-making.",
        "when": "Main emphasis is personal conduct, restraint, growth, or habit formation in a constructive sense.",
        "when_not": "Discipline is only about controlling women or asserting dominance — that goes to Authority and Submission. Self-help with manosphere undertones (alpha-male advice) goes to Relationship Tactics.",
        "rule": "Distinguishes constructive self-improvement from misogynistic self-help; both may use discipline language but point in different normative directions.",
        "why": "Captures the healthier self-help register. Maps to content codebook Q7 'improve themselves' and Q9 'self-discipline/fitness'.",
        "evidence": "4 audience, 4 content. Stable at 87% precision but low volume.",
        "creators": "Banky audience minor, Deyemi minor.",
    },
    {
        "code": "Marriage and Family",
        "bundle": "Cross-cutting (used in both progressive and regressive framings)",
        "definition": "Marriage, divorce, infidelity, husband/wife relations, fatherhood/motherhood, child support, family structure, household duty.",
        "when": "Marriage or family is the main object of discussion, not merely the setting where another gender norm appears.",
        "when_not": "When the primary claim is about submission, provider status, or sexual morality and marriage is incidental — use the more specific theme.",
        "rule": "Use as primary ONLY when the institution of marriage / family itself is the central topic. Apply the strict gating rule to avoid lazy defaults.",
        "why": "Most Nigeria audience comments and a large share of content are framed through family life. Coding this theme lets us compare progressive partnership / fatherhood narratives against regressive endurance / submission narratives within the same setting.",
        "evidence": "108 audience, 97 content — the largest theme. Stable at 95.6% precision but at risk of becoming a default.",
        "creators": "Agba (71 audience), Banky (27 audience, 42 content), all creators secondary.",
    },
    {
        "code": "Faith and Moral Repair",
        "bundle": "Cross-cutting",
        "definition": "Explicit faith, scripture, God, prayer, church, sin, spiritual testimony tied to masculinity, marriage, healing, or moral conduct.",
        "when": "Spiritual language is central to the gender / masculinity claim.",
        "when_not": "Generic morality or 'truth' statements without religious framing. Religion mentioned only as a passing reference or as cultural marker.",
        "rule": "Must include explicit spiritual vocabulary (God, prayer, scripture, sin, church, faith, testimony). Otherwise re-route to Marriage and Family or another theme.",
        "why": "Faith is culturally salient in the Nigeria corpus, especially Banky's testimony content and Agba's marriage advice. Without strict gating, this theme over-applies (44% precision in v2 audit). With strict gating, it picks up the specific religion-tied masculine framings worth tracking.",
        "evidence": "11 audience, 7 content. Flagged for tightening — strict rule expected to drop count to 5–7 audience.",
        "creators": "Agba (8 audience), Banky content, Shola minor.",
    },
]

UNCLEAR = {
    "code": "Unclear",
    "bundle": "Meta",
    "definition": "Low-signal, off-topic, or genuinely uncodable rows.",
    "when": "Rare cases where no real theme fits even after applying full decision rules.",
    "rule": "Cap usage at 5–7% of all rows. If higher, coding is too conservative.",
    "why": "A controlled parking spot. Replaces the prior 'OTHER' label, which was not in the codebook and broke aggregation.",
}

SCOPE = {
    "code": "Masculinity Identity",
    "bundle": "Scope marker (boolean column, not a primary theme)",
    "definition": "TRUE when the row directly discusses men, boys, manhood, or masculinity as a group, including male socialization or male identity. FALSE otherwise.",
    "when": "As a scope check on every row, never as a primary theme.",
    "rule": "If the row centers male experience as a group, set TRUE — but use a more specific code for the primary theme. Demoted from primary because at 211 mentions in v2 it functioned as an umbrella that swallowed more specific themes.",
    "why": "Acts as scope gate for the corpus. Lets analysts filter rows that explicitly invoke masculinity as a category, separate from the substantive theme being coded.",
}

AUX = [
    {
        "code": "Audience Stance",
        "scope": "Audience rows only",
        "definition": "What the commenter is doing in response to the source content.",
        "values": ["Support — backs the creator's point",
                   "Challenge — pushes back, disagrees",
                   "Mixed — partial agreement and partial pushback",
                   "Question — asks for clarification or extension",
                   "Testimony — first-person personal experience",
                   "Joke-casual — humor / light-hearted, no substantive position"],
        "why": "Topic theme alone cannot tell whether a Gender-Based Violence and Consent comment supports the anti-rape message or pushes back with whataboutism. Without stance, audience analysis is half-blind.",
    },
    {
        "code": "Sentiment",
        "scope": "Both audience and content (matches human codebook Q1)",
        "definition": "The overall valence of the text. Captures whether the comment expresses a positive, negative, or neutral disposition. Unclear is reserved for cases where valence cannot be determined (e.g., rhetorical questions or fragmentary text).",
        "values": ["Positive",
                   "Negative",
                   "Neutral",
                   "Unclear"],
        "why": "Mirrors the human audience codebook Q1 vocabulary so that LLM-coded sentiment is directly comparable to human-coded sentiment for inter-rater reliability checks. Earlier 4-class schemes that included a Mixed category were dropped because they produced low-precision splits and were not usable for the prior project's standard reporting.",
    },
    {
        "code": "Emotion Detection",
        "scope": "Both audience and content (mirrors human codebook Q2 vocabulary)",
        "definition": "The dominant discrete emotion expressed by the commenter or speaker. Built directly from the human codebook's Q2 'Primary emotional tone' option list to keep LLM-coded emotion comparable to human-coded emotion.",
        "values": ["Joy",
                   "Happiness",
                   "Surprise",
                   "Anger",
                   "Fear",
                   "Contempt",
                   "Sadness",
                   "Hope",
                   "Empathy",
                   "None of these"],
        "why": "Q2 of the human codebook bundles 'emotion' and 'tone' into a single field labeled 'Primary emotional tone'. The final codebook splits these into two distinct dimensions: Emotion Detection captures what the commenter feels (this field), and Tone captures how the message is delivered (next field). Splitting these axes is supported by classical work in social-media discourse analysis, where emotion (psychological state) and tone (rhetorical register) consistently behave as separable dimensions. Emotion Detection retains the Q2 vocabulary verbatim so the LLM output can be cross-validated against any future human pass that uses Q2 directly.",
    },
    {
        "code": "Tone",
        "scope": "Both audience and content (NEW dimension; not directly in human codebook)",
        "definition": "The rhetorical register of the comment — how the message is delivered, distinct from what the speaker feels. Captures stylistic posture toward the topic and the audience.",
        "values": ["Earnest — sincere, direct, no irony",
                   "Sarcastic — ironic, mocking, opposite-meaning",
                   "Hostile — aggressive, attacking, confrontational",
                   "Humorous — playful, joking, light-hearted",
                   "Empathetic — supportive, compassionate, validating",
                   "Authoritative — didactic, prescriptive, lecturing",
                   "Detached — neutral, observational, distant"],
        "why": "Tone is not separately captured anywhere in the human audience or content codebook — Q2 conflates emotion and tone. This is a meaningful gap: a comment expressing the emotion 'Anger' could be delivered in an earnest register, a sarcastic register, or a hostile register, and these distinctions matter substantively for the playbook (e.g., distinguishing sincere male-accountability advocacy from sarcastic deflection). The seven values are drawn from established discourse-analytic categories (Pennebaker LIWC analytical/casual axis; Berger 2014 marketing-tone framework; classical pragmatics on register). The list is deliberately compact (seven mutually exclusive categories) so the LLM and any future human coder can apply it consistently. Tone allows the analysis to separate, for example, an earnest Gender-Based Violence and Consent reply from a hostile one within the same theme bucket.",
    },
    {
        "code": "Normative Orientation",
        "scope": "Both audience and content",
        "definition": "Whether the text reinforces hierarchy / misogyny, challenges it, combines both, or is unclear.",
        "values": ["Progressive — challenges hierarchy or supports equality",
                   "Regressive — reinforces hierarchy or traditional roles",
                   "Mixed — both elements present",
                   "Unclear — cannot determine"],
        "why": "Separates topic from position. The same theme can appear in both progressive and regressive forms — Marriage and Family is the clearest example. This dimension makes the cross-tab coherent.",
    },
    {
        "code": "Rhetorical Strategy",
        "scope": "Both audience and content",
        "definition": "How the message persuades, not only what it says.",
        "values": ["Advice / Rule — prescriptive ('you should…')",
                   "Testimony — first-person personal experience",
                   "Moral Warning — 'this will end badly'",
                   "Ridicule — mockery, sarcasm, dismissal",
                   "Religious Appeal — scripture or doctrine-based",
                   "Whataboutism — deflection ('but what about X?')",
                   "Common-Sense Claim — 'everyone knows…', 'it's just reality'",
                   "Empathy and Solidarity — 'I see you / we stand with you'"],
        "why": "Maps onto content codebook Q10 (communication mode) and Q13 (justification). Surfaces the rhetorical fingerprint of each creator and the patterns of audience uptake.",
    },
    {
        "code": "Target of Claim",
        "scope": "Both audience and content",
        "definition": "Who or what the claim is directed at — the object of the speech act.",
        "values": ["Men / boys",
                   "Women / girls",
                   "Wives / girlfriends",
                   "Husbands / boyfriends",
                   "Feminists / modern women",
                   "Children / family",
                   "Institutions / law / society",
                   "Self / personal life",
                   "Mixed / unclear"],
        "why": "Disambiguates attacks-on-women from advice-to-men from defending-victims from personal narrative — distinctions that any single theme code cannot capture alone.",
    },
]

# ─── document build ─────────────────────────────────────────────────────────

doc = Document()

# Title block
t = doc.add_heading("Nigeria Manfluencer Project", level=0)
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("Final Coding Codebook")
r.bold = True; r.font.size = Pt(16)
sub2 = doc.add_paragraph()
sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub2.add_run("Norman Lear Center, USC Annenberg — Gates Foundation")
r.italic = True; r.font.size = Pt(11)
sub3 = doc.add_paragraph()
sub3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub3.add_run("Audience Reception and Content Analysis")
r.font.size = Pt(11)

doc.add_paragraph()

# 1. Purpose
add_heading(doc, "1. Purpose of this codebook", level=1)
para(doc, "This document defines the final coding scheme for the Nigeria portion of the Manfluencer Project. It is the canonical reference for both the LLM-assisted coding pipeline and any human validation pass. The themes apply to two coordinated tracks:")
bullet(doc, "Audience reception analysis: 417 manager-curated comments across four creators (Banky Wellington, Deyemi Okanlawon, Agba John Doe, Shola).", bold_lead="• Audience: ")
bullet(doc, "Content analysis: 310 coding-unit segments across six creators (the four above plus Wizarab and Ebuka Obi-Uchendu).", bold_lead="• Content: ")
para(doc, "The two tracks share an identical theme vocabulary so that creator framings and audience uptakes can be directly cross-tabulated. The audience track adds one extra auxiliary field — Audience Stance — to capture reception-specific behavior (support, challenge, testimony, deflection).", italic=False)

# 2. Methodology
add_heading(doc, "2. How this codebook was derived", level=1)
para(doc, "The final theme list emerged from four iterations of triangulation across the project's anchor documents and the data itself:")
bullet(doc, "Project scope memo (Gates Masculinity Scope - Streamlined.docx) — names the substantive themes the corpus is meant to capture: marriage, faith, provision, cheating, sexual morality, female sexuality, status, victim protection, gender grievance, sexual violence, male accountability, emotional openness, and male trauma.", bold_lead="Source 1 — ")
bullet(doc, "Nigeria selection memo (NIGERIA - Content and Audience Analysis Samples.docx) — provides per-creator framing notes that justify each theme's inclusion (e.g., 'victim protection' for Deyemi, 'marriage-market logic' for Agba, 'gender grievance' for Wizarab and Shola).", bold_lead="Source 2 — ")
bullet(doc, "Human content and audience codebooks (per-coder XLSX, audience: 41 columns / content: 33 columns) — formalize the question structure used by manual coders; the multi-select topic, problem, and solution items in those codebooks were cross-mapped against the LLM theme set to find gaps.", bold_lead="Source 3 — ")
bullet(doc, "Empirical data — 417 audience rows and 310 content rows from the locked Nigeria Audience Analysis Final and Nigeria Content Analysis Final workbooks. Candidate themes were tested against actual data via keyword and semantic probes; themes that fired on fewer than ~1% of rows (such as a separate 'emotional repression' or 'male friendship/community' code) were rejected because they cannot sustain statistical analysis.", bold_lead="Source 4 — ")
para(doc, "The list was then audited via a 15-pass validation procedure that checked row coverage, primary-theme textual support, top-3 agreement with an independent semantic probe, and per-theme precision. Themes flagged below 70% precision (Faith and Moral Repair, Relationship Tactics, Authority and Submission) were retained but given strict decision rules. Themes scoring above 90% precision (Marriage and Family, Provider and Status, Sexual Morality, Gender-Based Violence and Consent, Egalitarian Partnership) were retained as-is.")

# Why these names (style)
add_heading(doc, "3. Naming conventions", level=1)
para(doc, "Theme names are written in plain English Title Case for human readability. Underscored machine codes (such as 'AUTHORITY_SUBMISSION' or 'GBV_CONSENT') are only used inside the LLM coding pipeline and are not surfaced in deliverables. Each theme below shows the human-readable label as the canonical form.")

# 4. Final list
add_heading(doc, "4. Final theme list", level=1)
para(doc, "Thirteen primary themes plus one meta code (Unclear) and one scope marker (Masculinity Identity). All themes apply to both audience and content tracks.")

t = doc.add_table(rows=1, cols=2)
t.style = "Light Grid Accent 1"
hdr = t.rows[0].cells
hdr[0].text = "Theme"
hdr[1].text = "Bundle"
for c in hdr:
    for p in c.paragraphs:
        for r in p.runs: r.bold = True; r.font.size = Pt(10)
for th in THEMES:
    row = t.add_row().cells
    row[0].text = th["code"]
    row[1].text = th["bundle"]
    for c in row:
        for p in c.paragraphs:
            for r in p.runs: r.font.size = Pt(10)
row = t.add_row().cells
row[0].text = UNCLEAR["code"]; row[1].text = UNCLEAR["bundle"]
for c in row:
    for p in c.paragraphs:
        for r in p.runs: r.font.size = Pt(10); r.italic = True

doc.add_paragraph()
para(doc, "Plus the scope marker Masculinity Identity (TRUE/FALSE), described in Section 6.")

# 5. Theme-by-theme detail
add_heading(doc, "5. Theme definitions", level=1)
para(doc, "Each entry below specifies the theme's definition, decision rules, the project rationale, and the empirical basis. Decision rules govern when to use the theme as primary versus secondary versus not at all, and are designed to be applied consistently by both LLM coders and human validators.")

for i, th in enumerate(THEMES, 1):
    add_heading(doc, f"5.{i} {th['code']}", level=2)
    para(doc, f"Bundle: {th['bundle']}", italic=True, size=10)
    para(doc, "Definition", bold=True, space_after=2)
    para(doc, th["definition"], space_after=8)
    para(doc, "When to use as primary", bold=True, space_after=2)
    para(doc, th["when"], space_after=8)
    para(doc, "When NOT to use", bold=True, space_after=2)
    para(doc, th["when_not"], space_after=8)
    para(doc, "Decision rule", bold=True, space_after=2)
    para(doc, th["rule"], space_after=8)
    para(doc, "Why this theme is in the codebook", bold=True, space_after=2)
    para(doc, th["why"], space_after=8)
    para(doc, "Empirical basis", bold=True, space_after=2)
    para(doc, th["evidence"], space_after=8)
    para(doc, "Where it concentrates", bold=True, space_after=2)
    para(doc, th["creators"], space_after=12)

# Unclear
add_heading(doc, f"5.{len(THEMES)+1} {UNCLEAR['code']} (meta)", level=2)
para(doc, f"Bundle: {UNCLEAR['bundle']}", italic=True, size=10)
para(doc, "Definition", bold=True, space_after=2)
para(doc, UNCLEAR["definition"], space_after=8)
para(doc, "When to use", bold=True, space_after=2)
para(doc, UNCLEAR["when"], space_after=8)
para(doc, "Decision rule", bold=True, space_after=2)
para(doc, UNCLEAR["rule"], space_after=8)
para(doc, "Why it exists", bold=True, space_after=2)
para(doc, UNCLEAR["why"], space_after=12)

# 6. Scope marker
add_heading(doc, "6. Scope marker — Masculinity Identity", level=1)
para(doc, SCOPE["definition"], space_after=8)
para(doc, "When to use", bold=True, space_after=2)
para(doc, SCOPE["when"], space_after=8)
para(doc, "Decision rule", bold=True, space_after=2)
para(doc, SCOPE["rule"], space_after=8)
para(doc, "Why it is a scope marker, not a theme", bold=True, space_after=2)
para(doc, SCOPE["why"], space_after=12)

# 7. Auxiliary fields
add_heading(doc, "7. Auxiliary fields", level=1)
para(doc, "These fields supplement the primary and secondary theme codes. They capture how a claim is made (rhetorical strategy), whether it is reinforcing or challenging traditional norms (normative orientation), who the claim is directed at (target of claim), and — for audience rows only — what the commenter is doing relative to the source content (audience stance).")

for j, aux in enumerate(AUX, 1):
    add_heading(doc, f"7.{j} {aux['code']}", level=2)
    para(doc, f"Scope: {aux['scope']}", italic=True, size=10)
    para(doc, "Definition", bold=True, space_after=2)
    para(doc, aux["definition"], space_after=8)
    para(doc, "Values", bold=True, space_after=2)
    for v in aux["values"]:
        bullet(doc, v)
    para(doc, "Why this field is included", bold=True, space_after=2)
    para(doc, aux["why"], space_after=12)

# 8. Track-specific schemas
add_heading(doc, "8. Coding schemas by track", level=1)
para(doc, "Both tracks share the 13 primary themes, the Unclear meta code, and the Masculinity Identity scope marker. They differ only in the auxiliary set.")
para(doc, "Audience analysis schema (11 fields per row)", bold=True)
for line in [
    "Primary Theme — one of the 13 themes or Unclear",
    "Secondary Theme 1 — one of the 13 themes or blank",
    "Secondary Theme 2 — one of the 13 themes or blank",
    "Masculinity Identity — TRUE / FALSE",
    "Sentiment — Positive / Negative / Neutral / Unclear",
    "Emotion Detection — one of ten emotion values (Joy, Happiness, Surprise, Anger, Fear, Contempt, Sadness, Hope, Empathy, None of these)",
    "Tone — Earnest / Sarcastic / Hostile / Humorous / Empathetic / Authoritative / Detached",
    "Audience Stance — Support / Challenge / Mixed / Question / Testimony / Joke-casual",
    "Normative Orientation — Progressive / Regressive / Mixed / Unclear",
    "Rhetorical Strategy — one of eight",
    "Target of Claim — one of nine",
]: bullet(doc, line)

para(doc, "Content analysis schema (10 fields per row)", bold=True)
for line in [
    "Primary Theme — one of the 13 themes or Unclear",
    "Secondary Theme 1 — one of the 13 themes or blank",
    "Secondary Theme 2 — one of the 13 themes or blank",
    "Masculinity Identity — TRUE / FALSE",
    "Sentiment — Positive / Negative / Neutral / Unclear",
    "Emotion Detection — one of ten emotion values",
    "Tone — Earnest / Sarcastic / Hostile / Humorous / Empathetic / Authoritative / Detached",
    "Normative Orientation — Progressive / Regressive / Mixed / Unclear",
    "Rhetorical Strategy — one of eight",
    "Target of Claim — one of nine",
]: bullet(doc, line)

para(doc, "Operational LLM-coding schema (matches human codebook Q1–Q21h)", bold=True)
para(doc, "When coding is performed by the LLM pipeline (notebook: 'LLM Coding Notebook - Audience Analysis.ipynb'), the output spreadsheet additionally carries every Q1 through Q21h column from the human audience codebook so that LLM and human coding can be compared field-by-field. The four front summary columns (Themes, Sentiment, Emotion Detection, Tone) are concise rollups of the same underlying judgments captured in Q1, Q2, and the theme list. Inter-rater reliability between the LLM and any future human pass should therefore be computed on the Q-numbered columns where vocabularies match exactly.", space_after=8)

# 9. Decision rules — disambiguation
add_heading(doc, "9. Disambiguation rules", level=1)
para(doc, "The themes overlap in places. The following rules govern primary-theme assignment when more than one code seems to fit. These rules are listed in order of priority — the first applicable rule wins.")
disambig = [
    ("Violence trumps morality.", "If rape, consent, abuse, or victim/accuser framing is in scope, code Gender-Based Violence and Consent — even if sexual morality language is also present."),
    ("Hierarchy trumps family.", "If the text endorses, debates, or challenges hierarchy between men and women, code Authority and Submission — even if marriage is the setting."),
    ("Specificity trumps Marriage and Family.", "Marriage and Family is the most over-applied default. Use it only when the institution itself is the central topic. If the row is really about submission, sexual double standards, or partnership norms, use the more specific code."),
    ("Specificity trumps Masculinity Identity.", "Masculinity Identity is a scope marker. Set it TRUE alongside any specific theme. Never use it as primary."),
    ("Strict gating on Faith and Moral Repair.", "Requires explicit faith / scripture / God / prayer / sin language. Generic moral judgment without a religious frame goes to a different theme."),
    ("Tactics versus values.", "Relationship Tactics requires tactical advice or strategy. Aspirational or values-based relationship talk goes to Egalitarian Partnership or Marriage and Family."),
    ("Grievance versus victimhood.", "Gender Grievance frames women / feminists as the threat. Male Victimhood frames men as the harmed party. Both can co-occur as primary plus secondary."),
    ("Repression versus openness.", "Trauma and Mental Health captures the openness side of emotional discourse. The repression side ('men shouldn't show weakness') is captured by combining Trauma and Mental Health as primary with Normative Orientation = Regressive."),
]
for k, v in disambig:
    para(doc, k, bold=True, space_after=2)
    para(doc, v, space_after=8)

# 10. Version history
add_heading(doc, "10. Version history", level=1)
versions = [
    ("v1 — Initial controlled vocabulary (April–May 2026)",
     "15 codes derived directly from the project scope memo. Used in the first exploratory analysis pass on 417 audience and 310 content rows. Surfaced gaps around male victimhood and male accountability framing."),
    ("v2 — 15-Pass Audit revision",
     "Independent ChatGPT audit of the v1 coded data flagged 21 OTHER assignments and three over-applied codes (FAITH, RELATIONSHIP_TACTICS, AUTHORITY_SUBMISSION). Reduced to 12 primary themes plus OTHER."),
    ("v3 — Gap analysis revision",
     "Cross-checked v2 against the human content and audience codebooks. Added Male Victimhood and Male Accountability as standalone themes (previously bundled into Gender Grievance and Gender-Based Violence and Consent respectively). Demoted Masculinity Identity from primary theme to scope marker. Added Target of Claim as a fourth auxiliary field. Replaced OTHER with Unclear and capped its usage at 5–7%."),
    ("v4 — Empirical re-validation",
     "Tested two further candidate themes (Emotional Repression, Male Friendship and Community) against the audience and content data. Both fired on under 1% of rows after manual review of regex candidates, and were rejected as primary themes — captured instead through Normative Orientation and secondary tags. The codebook is now stable at 13 primary themes plus Unclear plus Masculinity Identity scope marker plus four auxiliary fields."),
    ("v5 — Kenya validation",
     "Validated the v4 13-theme codebook against the Kenya audience corpus (412 comments × 4 creators: Andrew Kibe, Onyango Otieno / Rixpoet, Eddy Kimani, Eric Amunga / Amerix). 20 keyword probes were applied: the 13 codebook themes plus 7 Kenya-specific candidates (Polygamy, Sheng/Swahili gendered terms, Tribalism/regional, Politics/governance, FGM, Fitness/Hustle, Education/skill). All 7 Kenya-specific candidates fired at under 1% of rows; none warranted a new theme. Distribution shifts versus Nigeria are findings rather than codebook flaws — Faith and Moral Repair fires more in Kenya (Christian/Muslim faith vocabulary common in Amerix's regressive register and Rixpoet's healing register), while Gender Grievance fires less. Same prompt and same 13 + Unclear vocabulary applied to a 200-comment stratified sample (50 per creator) yielded clean per-creator profiles: Trauma and Mental Health is Rixpoet-exclusive, Authority and Submission is Amerix-anchored, Self-Discipline is Eddy-anchored. The codebook is therefore confirmed as country-cross-cutting — same theme list applies to Nigeria and Kenya audience and content tracks."),
    ("v6 — Audit fixes (current)",
     "External audit on the v4-coded Nigeria xlsx flagged eight structural and logic issues: Themes column was a combined string instead of separate primary/secondary cells; Q16 (opinion reinforced) was over-coded at 74% of rows; Q20 (corrects something) and Q18 (shares info) were stuck at 0%; Q4 had 49 internal contradictions with Q5/Q6; Q11 was empty on 89% of Challenging rows; Q21 had 10 internal contradictions with sub-fields; theme codes occasionally applied regressive frames to comments critiquing those frames. Fixes applied in v6: split Themes into Primary Theme + Secondary Theme 1 + Secondary Theme 2; added Masculinity Identity, Normative Orientation, and Target of Claim as standalone columns; tightened the prompt with eight named CRITICAL CODING RULES (RULE 1 — themes reflect what the COMMENTER advocates, not the source post; RULE 2 — Q4 must be consistent with Q5/Q6; RULE 3 — Q11 required when Q8 = Challenging; RULE 4 — Q16 high bar, default No; RULE 5 — Q18 covers personal info as well as facts/links; RULE 6 — Q20 covers explicit correction language; RULE 7 — Q21 must be Yes if any of Q21a/c/e/g is Yes; RULE 8 — separate the commenter from the source creator); the validator auto-fixes Q4 when Q5/Q6 carry sentiment values, auto-fixes Q21 when any sub-field is Yes, and locks Q1 = Sentiment and Q2 = Emotion Detection. Re-running both Nigeria and Kenya at gpt-4o-mini temperature 0 reduced all eight issue rates to acceptable bounds: Q4 inconsistencies 0/200 in both; Q11 missing 4/145 in Nigeria and 0/86 in Kenya; Q21 inconsistencies 0; Q16 = 0 (empirically defensible — no explicit reinforcement language found in either 200-row sample after a phrase scan); Q20 firing on 51 Nigeria rows (25.5%) and 5 Kenya rows; Q18 firing on 9 Nigeria rows and 48 Kenya rows. Theme coding for Challenging rows is now defensible — spot-check confirms most challenger comments who carry regressive primary themes are themselves advancing those themes (e.g., a comment that pushes back on a creator's specific advice while still expressing male-victimhood grievance)."),
]
for label, body in versions:
    para(doc, label, bold=True, space_after=2)
    para(doc, body, space_after=8)

# 11. Application notes
add_heading(doc, "11. Application notes", level=1)
para(doc, "How to read a coded row.", bold=True, space_after=2)
para(doc, "A single row's full coding has the shape: Primary Theme + up to two Secondary Themes + Masculinity Identity scope flag + Normative Orientation + Rhetorical Strategy + Target of Claim, plus Audience Stance for audience rows. The primary theme answers what the comment is mainly about; secondaries capture adjacent frames; the auxiliaries describe how the claim is made and where it is aimed.", space_after=8)
para(doc, "How to handle close calls.", bold=True, space_after=2)
para(doc, "When two themes seem equally plausible, apply Section 9's disambiguation rules in order. If still tied, prefer the more specific theme; only fall back to Marriage and Family or Masculinity Identity scope when no more specific theme fits.", space_after=8)
para(doc, "How to interpret the cross-track comparison.", bold=True, space_after=2)
para(doc, "Because audience and content use the same theme set, the central comparative metrics are: (a) echo ratio — proportion of audience comments that match the primary theme of the source content, (b) theme injection — themes that audiences add that were not in source content, (c) theme attrition — themes creators emphasize that audiences ignore, and (d) friction — audience comments coded Challenge stance with opposite Normative Orientation. These metrics are computable directly from the schemas above.", space_after=8)
para(doc, "Reproducibility.", bold=True, space_after=2)
para(doc, "All LLM coding runs are cacheable by SHA-1 hash of input text. Re-running the pipeline with the same seed and same prompt is deterministic at temperature 0. Multi-temperature stability runs (15 to 25 iterations across mixed temperatures) are recommended for any contested row but are not required for the headline distributions.", space_after=8)

# Footer
para(doc, "")
para(doc, "Document prepared for the Norman Lear Center, USC Annenberg / Gates Foundation Manfluencer Project. Theme set finalized following four iterations of validation against the project scope, the human codebooks, and the empirical Nigeria corpus.", italic=True, size=10)

doc.save(OUT)
print(f"wrote {OUT}")
print(f"size: {OUT.stat().st_size:,} bytes")
