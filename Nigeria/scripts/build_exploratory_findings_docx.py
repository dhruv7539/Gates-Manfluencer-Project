"""
Build the bullet-proofed LLM Exploratory Findings docx.

Goals:
- Fix the 6 internal inconsistencies identified in the v1 draft.
- Address Ksenia's 4 in-doc comments + Slack thread:
    * normalization / oversample acknowledgment at the headline figure
    * denominators stated everywhere
    * how figures should be interpreted (in the doc, not the deck)
    * scope-of-document clarifier (exploratory only; engagement metrics deferred)
- Strip AI-flavored prose (em-dash chains, "this is the X door", repetitive scope-quote preambles, etc.)
- Keep the doc readable in Google Docs / Word with simple Arial styling.
"""

from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUT_PATH = Path(
    "Codebooks/LLM Codebook/LLM Exploratory Analysis.docx"
)

# ---------------------------------------------------------------------------
# Styling helpers
# ---------------------------------------------------------------------------

ARIAL = "Arial"
COLOR_HEAD = RGBColor(0x1F, 0x2A, 0x44)   # dark navy
COLOR_BODY = RGBColor(0x14, 0x14, 0x14)   # near-black
COLOR_MUTED = RGBColor(0x55, 0x55, 0x55)
COLOR_RULE_BG = "F2F4F7"                  # light grey for callout cells


def _set_run(run, *, size=10.5, bold=False, italic=False, color=COLOR_BODY):
    run.font.name = ARIAL
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), ARIAL)
    rfonts.set(qn("w:hAnsi"), ARIAL)
    rfonts.set(qn("w:cs"), ARIAL)
    rfonts.set(qn("w:eastAsia"), ARIAL)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color


def _add_para(doc, text="", *, size=10.5, bold=False, italic=False,
              color=COLOR_BODY, align=None, space_after=4):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    if text:
        run = p.add_run(text)
        _set_run(run, size=size, bold=bold, italic=italic, color=color)
    return p


def _add_heading(doc, text, *, level=1):
    sizes = {1: 16, 2: 12.5, 3: 11}
    size = sizes.get(level, 11)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    _set_run(run, size=size, bold=True, color=COLOR_HEAD)
    return p


def _set_cell_shade(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def _set_cell_borders(cell, color="BFBFBF", size="4"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), size)
        b.set(qn("w:color"), color)
        borders.append(b)
    tc_pr.append(borders)


def _write_cell(cell, text, *, bold=False, size=10, color=COLOR_BODY,
                align=None, shade=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    _set_run(run, size=size, bold=bold, color=color)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    if shade is not None:
        _set_cell_shade(cell, shade)
    _set_cell_borders(cell)


def _add_table(doc, header, rows, *, col_widths=None,
               header_shade="1F2A44", header_color=RGBColor(0xFF, 0xFF, 0xFF),
               first_col_bold=False):
    table = doc.add_table(rows=1 + len(rows), cols=len(header))
    table.autofit = False
    table.allow_autofit = False
    # widths
    if col_widths is not None:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    # header
    for i, h in enumerate(header):
        _write_cell(
            table.rows[0].cells[i], h,
            bold=True, size=9.5, color=header_color,
            align=WD_ALIGN_PARAGRAPH.LEFT, shade=header_shade,
        )
    # rows
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            _write_cell(
                table.rows[r + 1].cells[c], str(val),
                bold=(first_col_bold and c == 0),
                size=10,
                align=WD_ALIGN_PARAGRAPH.LEFT,
                shade=("FAFAFA" if r % 2 == 1 else None),
            )
    # spacing
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def _add_callout(doc, title, body, *, shade=COLOR_RULE_BG):
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    cell = table.rows[0].cells[0]
    cell.text = ""
    _set_cell_shade(cell, shade)
    _set_cell_borders(cell, color="D0D5DD")

    p1 = cell.paragraphs[0]
    p1.paragraph_format.space_after = Pt(2)
    run = p1.add_run(title)
    _set_run(run, size=10, bold=True, color=COLOR_HEAD)
    for line in body:
        p = cell.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(line)
        _set_run(run, size=10, color=COLOR_BODY)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def _add_bullets(doc, items, *, indent=0.4):
    for item in items:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(indent)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run("• " + item)
        _set_run(run, size=10.5)


# ---------------------------------------------------------------------------
# Document content
# ---------------------------------------------------------------------------

def build():
    doc = Document()
    # tighten margins
    for section in doc.sections:
        section.top_margin = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        section.left_margin = Cm(1.8)
        section.right_margin = Cm(1.8)

    # ---------- Title block ----------
    _add_para(
        doc, "Phase 1 Interim Findings",
        size=20, bold=True, color=COLOR_HEAD, space_after=2,
    )
    _add_para(
        doc, "Gates Masculinity Project · Components 2 and 3",
        size=12.5, color=COLOR_MUTED, space_after=2,
    )
    _add_para(
        doc, "Content analysis and audience reception · Kenya × Nigeria",
        size=11, color=COLOR_MUTED, space_after=12,
    )

    # ---------- Scope-of-document clarifier (NEW) ----------
    _add_callout(
        doc,
        "Scope of this document",
        [
            "This deliverable reports the exploratory LLM coding pass that "
            "covers Component 2 (Content analysis) and Component 3 "
            "(Audience reception analysis). The focused, hypothesis-driven "
            "follow-up pass will be circulated separately.",
            "Engagement-metric correlation (view counts, impressions, "
            "shares × content variables) is reported in the final deck. "
            "View counts are unavailable on some platforms in the sample; "
            "metrics are used wherever available and gaps are noted in "
            "Section 24.",
            "Findings should be read at the creator, archetype, and "
            "snippet/comment level. Corpus-level percentages are reported "
            "with their denominators in every table and are not "
            "population estimates (see Section 25 on sampling).",
        ],
    )

    # ---------- Sampling caveat right at the headline ----------
    _add_heading(doc, "Headline figures", level=2)
    _add_para(
        doc,
        "All figures below are shares of coded units within the curated "
        "Phase 1 sample. The sample was deliberately constructed to "
        "include creators across the regressive–progressive spectrum so "
        "that both ends would be observable; corpus shares therefore "
        "describe what is in the corpus, not what is typical of "
        "masculinity discourse in either country. Per-creator figures in "
        "Section 2 are the primary unit of analysis.",
    )

    _add_table(
        doc,
        header=["Metric", "Kenya", "Nigeria"],
        rows=[
            ["Audience comments coded (n)", "412", "417"],
            ["Content snippets coded (n)",  "394", "381"],
            ["Regressive share — % of country content snippets",
             "43% (171 / 394)", "43% (165 / 381)"],
            ["Progressive share — % of country content snippets",
             "35% (138 / 394)", "31% (117 / 381)"],
            ["Mixed / unclear — % of country content snippets",
             "22% (85 / 394)",  "26% (99 / 381)"],
        ],
        col_widths=[8.5, 4, 4],
        first_col_bold=True,
    )
    _add_para(doc,
        "Read every percentage in this document as \"share within the "
        "curated Phase 1 sample\". Because creators were selected to "
        "span both ends of the regressive ↔ progressive axis, the "
        "headline 43% / 35% / 31% figures reflect the composition of "
        "the corpus, not a population estimate for either country. "
        "Per-creator figures (Section 2) are the primary unit of "
        "analysis.",
        size=10, italic=True, color=COLOR_MUTED,
    )

    _add_para(
        doc,
        "Total coded rows: 1,604 (829 audience comments and 775 content "
        "snippets across 11 selected creators — 5 Kenya, 6 Nigeria). "
        "Sentiment was coded on a four-valued scale "
        "(positive · negative · neutral · unclear) at temperature 0 "
        "with cached, deterministic outputs (see Section 24).",
    )

    # ---------- Executive summary ----------
    _add_heading(doc, "Executive summary", level=2)

    _add_para(doc,
        "1.  Two distinct masculinity registers. The most common Kenyan "
        "masculinity narrative is \"men are disadvantaged victims\" "
        "(92 / 394 content snippets = 23%); it is the top narrative for "
        "Rixpoet and Eddy Kimani (2 of 5 Kenyan creators). The most "
        "common Nigerian narrative is \"men should provide and succeed\" "
        "(79 / 381 = 21%), which is the top narrative for three of six "
        "Nigerian creators (Ebuka Obi-Uchendu, Agba John Doe, Wizarab). "
        "The two registers — grievance and provision — function as "
        "distinct entry points into the same ecosystem.",
    )
    _add_para(doc,
        "2.  Regressive content concentrates in a small set of creators "
        "within this sample. Kibe and Amerix together account for 157 "
        "of the 171 Kenyan regressive snippets (92%). In Nigeria, Agba "
        "John Doe, Shola, and Wizarab carry the bulk of regressive "
        "content; Banky Wellington, Deyemi Okanlawon, and Ebuka "
        "Obi-Uchendu code as progressive-leaning. Per-creator "
        "orientation is in Section 11.",
    )
    _add_para(doc,
        "3.  Solutions diagnosis is thin. 44% of Kenyan (172 / 394) and "
        "59% of Nigerian (226 / 381) content snippets propose no clear "
        "solution. Where solutions do appear, the leading non-empty "
        "categories are emotional growth and healing (Kenya: 102 "
        "snippets) and family responsibility (Nigeria: 54 snippets) — "
        "both compatible with counter-narrative framing.",
    )
    _add_para(doc,
        "4.  Audience response patterns diverge by country in this "
        "sample. Kenyan audiences validate more than they push back "
        "(validation 155, resistance / pushback 100). Nigerian "
        "audiences invert that pattern (pushback 150, validation 113). "
        "The Nigerian comment threads in this sample appear closer to "
        "organised counter-publics than to uniform echo chambers; this "
        "is a sample-level observation and is not generalised beyond "
        "the curated set.",
    )
    _add_para(doc,
        "5.  Within this sample, the healing / testimony register "
        "appears to carry the clearest counter-narrative signature. "
        "Under Rixpoet, audience comments code 88% supports and a high "
        "share of extends-the-message uptake. Banky Wellington shows "
        "the same pattern in Nigeria. This is also the only register "
        "in the sample that produces both learning / reflection and "
        "personal disclosure at meaningful rates.",
    )
    _add_para(doc,
        "6.  Audiences open a violence/safety frame that creators do not "
        "lead with. In Nigeria, violence/safety appears in 39 audience "
        "comments but only 3 content snippets — almost all under Deyemi "
        "Okanlawon. This is a frame the audience is asking for; creators "
        "in this sample are not supplying it.",
    )
    _add_para(doc,
        "7.  Hostile sexism is concentrated; benevolent sexism is the "
        "larger ambient channel. Hostile or dehumanising sexism appears "
        "in roughly 20% of regressive snippets in each country and "
        "concentrates in the same regressive-dominant creators. "
        "Benevolent or traditional sexism (\"submit\", \"headship\", "
        "\"real man provides\") is the wider register through which "
        "gender hierarchy is reproduced as common-sense.",
    )

    _add_callout(
        doc,
        "How this report maps to the Scope of Work",
        [
            "Core topics → §3 · Masculinity narratives → §4 · Themes and "
            "framing → §5 · Problems and solutions → §6.",
            "Sentiments → §7 · Rhetorical strategies → §8 · Argument "
            "types → §9 · Misogyny and sexism → §10 · Regressive ↔ "
            "progressive orientation → §11.",
            "Content cluster and entry-point map → §12–13 · Audience "
            "sentiment, emotion, perceived impact → §14–19.",
            "Linked content–audience analysis → §20–21 · Playbook "
            "implications (Component 4 bridge) → §22.",
            "Definitions and computation for every coded variable → §23. "
            "Methodology and limitations → §24–25.",
        ],
    )

    # ====================================================================
    # PART A · APPROACH
    # ====================================================================
    _add_heading(doc, "Part A · Scope and approach", level=2)

    _add_heading(doc, "1.  What this report covers", level=3)
    _add_para(doc,
        "The Scope of Work defines four components: (1) landscape "
        "analysis, (2) content analysis, (3) audience reception "
        "analysis, (4) playbook development. This report is the "
        "exploratory pass for Components 2 and 3, with explicit hooks "
        "into Component 4 in Section 22. Content analysis surfaces core "
        "topics, themes, narratives, sentiments, and rhetorical "
        "strategies; audience reception analysis surfaces sentiment, "
        "emotions, and the perceived impact of the content as reported "
        "by engaged audiences. Each scope-named output is a numbered "
        "section below."
    )

    _add_heading(doc, "2.  Sample composition", level=3)
    _add_para(doc,
        "Eleven masculinity-associated creators across Kenya and "
        "Nigeria, drawn from the manosphere-adjacent ecosystem "
        "identified in the landscape analysis. The sample sits within "
        "the Scope's 3–5-per-country guideline and 1,600-comment "
        "budget. Creators were chosen so that both ends of the "
        "regressive–progressive spectrum are visible in the data; this "
        "shapes how the corpus-level percentages should be read "
        "(see Section 25)."
    )

    _add_table(doc,
        header=["Creator", "Country", "Snippets (n)", "Comments (n)",
                "Reg. %", "Prog. %"],
        rows=[
            ["Onyango Otieno (Rixpoet)", "Kenya", "100", "92", "9%",  "49%"],
            ["Eric Amunga (Amerix)",     "Kenya", "97",  "140","86%", "12%"],
            ["Andrew Kibe",              "Kenya", "91",  "70", "81%", "13%"],
            ["Eddy Kimani",              "Kenya", "65",  "110","2%",  "71%"],
            ["Philip Karanja",           "Kenya", "41",  "0",  "10%", "46%"],
            ["Ebuka Obi-Uchendu",        "Nigeria","83", "0",  "16%", "34%"],
            ["Agba John Doe",            "Nigeria","77", "161","78%", "10%"],
            ["Banky Wellington",         "Nigeria","76", "110","8%",  "64%"],
            ["Shola",                    "Nigeria","72", "66", "78%", "14%"],
            ["Wizarab",                  "Nigeria","49", "0",  "59%", "14%"],
            ["Deyemi Okanlawon",         "Nigeria","24", "80", "4%",  "62%"],
        ],
        col_widths=[6.6, 2.4, 2.5, 2.8, 2.0, 2.0],
        first_col_bold=True,
    )
    _add_para(doc,
        "Reg. % and Prog. % are shares of that creator's content "
        "snippets coded as regressive or progressive respectively; the "
        "complement (24 percentage points on average) is the mixed or "
        "unclear share. Audience comments are not coded for "
        "regressive/progressive orientation; that axis applies to "
        "content snippets only.",
        size=9.5, color=COLOR_MUTED,
    )

    # ====================================================================
    # PART B · CONTENT FINDINGS
    # ====================================================================
    _add_heading(doc, "Part B · Content findings (Scope Component 2)", level=2)

    _add_heading(doc, "3.  Core topics", level=3)
    _add_para(doc,
        "Kenya's content centres on mental health (31% of content "
        "snippets); Nigeria's on dating and relationships (38%). The "
        "two ecosystems pivot around different topic doors."
    )
    _add_table(doc,
        header=["Top topics — share of content snippets",
                "Kenya (n=394)", "Nigeria (n=381)"],
        rows=[
            ["Mental health",          "31%", "4%"],
            ["Dating / relationships", "16%", "38%"],
            ["Gender debate",          "14%", "13%"],
            ["Marriage / family",      "—",   "16%"],
            ["Fatherhood / parenting", "13%", "—"],
            ["Money / status",         "—",   "12%"],
        ],
        col_widths=[7.5, 4.5, 4.5],
        first_col_bold=True,
    )

    _add_heading(doc, "4.  Masculinity narratives", level=3)
    _add_para(doc,
        "Each snippet is coded for the underlying masculinity model it "
        "advances (or \"none\" if the snippet is descriptive rather than "
        "prescriptive). At the corpus level:"
    )
    _add_bullets(doc, [
        "Kenya — the most frequent narrative is \"men are disadvantaged "
        "victims\" (92 snippets); it is the top narrative for Rixpoet "
        "and Eddy Kimani. \"Men should provide and succeed\" tops Kibe. "
        "\"Men should protect women and children\" tops Amerix and "
        "Philip Karanja. The progressive counter-narrative \"men should "
        "be emotionally open\" accounts for 41 snippets, concentrated "
        "under Rixpoet and Eddy.",
        "Nigeria — \"men should provide and succeed\" is the most "
        "frequent narrative (79 snippets) and tops Ebuka, Agba John "
        "Doe, and Wizarab. \"Men should protect women and children\" "
        "tops Banky Wellington. \"Men are disadvantaged victims\" tops "
        "Shola. \"Men should improve themselves\" tops Deyemi. The "
        "progressive counter-narrative \"men should be equal partners\" "
        "accounts for 40 snippets and is concentrated under Banky and "
        "Deyemi.",
    ])
    _add_callout(doc,
        "Cross-country narrative contrast",
        [
            "Vulnerability counter-narrative (\"emotionally open\") is "
            "alive in Kenya (41 snippets) and thin in Nigeria (14).",
            "Equality counter-narrative (\"equal partners\") is alive in "
            "Nigeria (40 snippets) and nearly absent in Kenya (4).",
        ],
    )

    _add_heading(doc, "5.  Themes and framing", level=3)
    _add_para(doc,
        "Each snippet and comment is coded for its dominant frame. The "
        "table below reports the top frames per country on the content "
        "side and the audience side; \"content\" and \"audience\" are "
        "counts in their respective corpora and are not directly "
        "comparable without their denominators."
    )
    _add_table(doc,
        header=["Frame", "KE content (n=394)", "KE audience (n=412)",
                "NG content (n=381)", "NG audience (n=417)"],
        rows=[
            ["Traditional gender order", "79", "58",  "61", "111"],
            ["Trauma healing",            "59", "62",  "5",  "1"],
            ["Self-improvement",          "51", "70",  "—",  "—"],
            ["Provider pressure",         "18", "59",  "—",  "—"],
            ["Female blame",              "36", "48",  "75", "82"],
            ["Violence / safety",         "—",  "—",   "3",  "39"],
        ],
        col_widths=[5.2, 3.2, 3.2, 3.2, 3.2],
        first_col_bold=True,
    )
    _add_bullets(doc, [
        "Kenya — content and audience track each other on traditional "
        "gender order and trauma healing. The audience opens "
        "self-improvement (51 → 70) and provider pressure (18 → 59) "
        "more often than the content does.",
        "Nigeria — the audience inflates traditional gender order "
        "(61 → 111) and female blame (75 → 82). Violence / safety "
        "appears in 39 audience comments under only 3 content snippets "
        "(see §20).",
    ])

    _add_heading(doc, "6.  Problems and solutions diagnosed", level=3)
    _add_para(doc,
        "Each snippet is coded for the problem it identifies (what is "
        "wrong) and the solution it prescribes (if any)."
    )
    _add_table(doc,
        header=["", "Kenya (n=394)", "Nigeria (n=381)"],
        rows=[
            ["Top problem", "Men's behaviour; mental health is unusually "
                            "high (66)",
                            "Men's behaviour; women / feminism; mental "
                            "health nearly absent (3)"],
            ["Top solution", "Emotional growth and healing (102)",
                             "Family responsibility (54); emotional "
                             "growth a distant second (36)"],
            ["No-clear-solution share", "172 / 394 = 44%",
                                        "226 / 381 = 59%"],
        ],
        col_widths=[4.5, 6.5, 6.5],
        first_col_bold=True,
    )

    _add_heading(doc, "7.  Three-axis sentiment", level=3)
    _add_para(doc,
        "Sentiment is coded separately for three targets within each "
        "content snippet: men, women, and traditional gender norms. "
        "Each row of each table below sums to the country's content-"
        "snippet n. Audience-side sentiment is reported in Section 14."
    )
    _add_para(doc, "Kenya · content snippets · n = 394", italic=True, size=10)
    _add_table(doc,
        header=["Target", "Positive", "Negative", "Neutral", "Unclear",
                "Row n"],
        rows=[
            ["Men",   "181", "98",  "115", "0", "394"],
            ["Women", "27",  "127", "234", "6", "394"],
            ["Norms", "65",  "153", "174", "2", "394"],
        ],
        col_widths=[3.5, 2.8, 2.8, 2.8, 2.5, 2.5],
        first_col_bold=True,
    )
    _add_para(doc, "Nigeria · content snippets · n = 381", italic=True, size=10)
    _add_table(doc,
        header=["Target", "Positive", "Negative", "Neutral", "Unclear",
                "Row n"],
        rows=[
            ["Men",   "128", "67",  "186", "0", "381"],
            ["Women", "48",  "137", "195", "1", "381"],
            ["Norms", "51",  "183", "147", "0", "381"],
        ],
        col_widths=[3.5, 2.8, 2.8, 2.8, 2.5, 2.5],
        first_col_bold=True,
    )
    _add_bullets(doc, [
        "Sentiment toward women is mostly neutral in both corpora. "
        "Content rarely takes an explicit positive or negative position "
        "on women; when it does, the modal direction is negative.",
        "Sentiment toward traditional norms is more polarised. Nigerian "
        "content is more critical of those norms (48% negative) than "
        "Kenyan content (39%).",
    ])

    _add_heading(doc, "8.  Rhetorical strategies", level=3)
    _add_para(doc,
        "Each snippet is coded for its dominant rhetorical mode "
        "(commentary, advice, debate, personal story, testimony, "
        "motivational speech, religious/moral teaching, humour, "
        "warning, or news/facts)."
    )
    _add_bullets(doc, [
        "Kenya — testimony and personal story concentrate under Rixpoet "
        "and Eddy (94 personal-story and 21 testimony snippets combined). "
        "Advice and instruction concentrate under Amerix.",
        "Nigeria — debate / argument is highest under Deyemi; "
        "commentary / opinion dominates the rest (237 snippets overall).",
    ])
    _add_callout(doc,
        "Why this matters",
        [
            "The testimony + personal-story register in Kenya correlates "
            "with the highest rates of audience personal disclosure and "
            "learning / reflection (see §16). It is the strongest single "
            "rhetorical signature of progressive content in this sample.",
        ],
    )

    _add_heading(doc, "9.  Argument types", level=3)
    _add_para(doc,
        "Generalisation dominates both corpora. Personal experience is "
        "the second-largest argument type in Kenyan content (driven by "
        "Rixpoet) but is overtaken by advice / instruction in Nigeria."
    )
    _add_table(doc,
        header=["Argument type (content)", "Kenya (n=394)", "Nigeria (n=381)"],
        rows=[
            ["Generalisation",       "169", "201"],
            ["Personal experience",  "106", "—"],
            ["Advice / instruction", "—",   "47"],
        ],
        col_widths=[7.5, 4.5, 4.5],
        first_col_bold=True,
    )

    _add_heading(doc, "10.  Misogyny and sexism (content side)", level=3)
    _add_para(doc,
        "Hostile and dehumanising sexism is rare in absolute volume but "
        "tightly concentrated. Benevolent / traditional sexism — the "
        "\"submit\", \"headship\", \"real man provides\" register — is "
        "the broader vehicle through which gender hierarchy is "
        "reproduced as common-sense."
    )
    _add_bullets(doc, [
        "Kenya — Amerix is the most hostile-sexism-heavy creator; "
        "Rixpoet and Eddy code 0%.",
        "Nigeria — Agba John Doe, Shola, and Wizarab carry most of the "
        "sexism load; Ebuka, Deyemi, and Banky code near zero.",
    ])

    _add_heading(doc, "11.  Regressive ↔ progressive orientation by creator",
                 level=3)
    _add_para(doc,
        "Each snippet is coded along a regressive ↔ progressive axis. "
        "Regressive: reinforces hierarchy, domination, female "
        "submission, misogyny, rigid provider roles, anti-feminism, or "
        "male grievance. Progressive: promotes equality, "
        "accountability, emotional openness, anti-violence, healing, or "
        "partnership. Snippets that do not fit either are coded mixed "
        "or unclear."
    )
    _add_bullets(doc, [
        "Kenya — Kibe and Amerix sit at ~85% regressive. Eddy and "
        "Rixpoet are progressive-leaning. Philip's small sample (41) "
        "skews progressive but the n is too small for a stable read.",
        "Nigeria — Agba John Doe, Shola, and Wizarab are regressive-"
        "dominant. Banky and Deyemi are progressive-dominant. Ebuka is "
        "mixed-leaning-progressive.",
    ])

    # ====================================================================
    # PART C · CONTENT CLUSTERS
    # ====================================================================
    _add_heading(doc,
                 "Part C · Content clusters and entry points", level=2)

    _add_heading(doc, "12.  Per-creator profile", level=3)
    _add_para(doc,
        "Each row is one creator's dominant cluster across the "
        "scope-named outputs. Read each row as the creator's profile, "
        "not as a ranking against other creators. Sex % is the share of "
        "that creator's content snippets coded as carrying any positive "
        "level of hostile, dehumanising, or benevolent/traditional "
        "sexism. Sup % and Opp % are shares of that creator's audience "
        "comments coded as supporting or opposing the original post."
    )
    _add_para(doc, "Kenya", italic=True, size=10)
    _add_table(doc,
        header=["Creator", "Top topic", "Top narrative", "Top frame",
                "Top rhetoric", "Reg %", "Prog %", "Sex %",
                "Sup %", "Opp %"],
        rows=[
            ["Rixpoet", "Mental health", "Men are disadvantaged victims",
             "Trauma healing", "Personal story",
             "9%", "49%", "0%", "88%", "10%"],
            ["Amerix", "Dating / relationships",
             "Men should protect women / children",
             "Traditional gender order", "Advice / instruction",
             "86%", "12%", "74%", "41%", "55%"],
            ["Kibe", "Marriage / family",
             "Men should provide / succeed",
             "Traditional gender order", "Commentary / opinion",
             "81%", "13%", "47%", "44%", "46%"],
            ["Eddy", "Mental health",
             "Men are disadvantaged victims", "Self-improvement",
             "Commentary / opinion", "2%", "71%", "0%", "64%", "32%"],
            ["Philip", "Violence / safety",
             "Men should protect women / children",
             "Fatherhood responsibility", "Commentary / opinion",
             "10%", "46%", "2%", "—", "—"],
        ],
        col_widths=[2.6, 2.6, 3.0, 2.6, 2.4, 1.2, 1.2, 1.2, 1.2, 1.2],
        first_col_bold=True,
    )
    _add_para(doc, "Nigeria", italic=True, size=10)
    _add_table(doc,
        header=["Creator", "Top topic", "Top narrative", "Top frame",
                "Top rhetoric", "Reg %", "Prog %", "Sex %",
                "Sup %", "Opp %"],
        rows=[
            ["Ebuka", "Dating / relationships",
             "Men should provide / succeed",
             "Provider pressure", "Commentary / opinion",
             "16%", "34%", "6%", "—", "—"],
            ["Agba JD", "Marriage / family",
             "Men should provide / succeed",
             "Traditional gender order", "Commentary / opinion",
             "78%", "10%", "64%", "53%", "41%"],
            ["Banky W.", "Fatherhood / parenting",
             "Men should protect women / children",
             "Fatherhood responsibility", "Commentary / opinion",
             "8%", "64%", "3%", "69%", "17%"],
            ["Shola", "Dating / relationships",
             "Men are disadvantaged victims",
             "Female blame", "Commentary / opinion",
             "78%", "14%", "46%", "52%", "48%"],
            ["Wizarab", "Dating / relationships",
             "Men should provide / succeed",
             "Female blame", "Commentary / opinion",
             "59%", "14%", "41%", "—", "—"],
            ["Deyemi", "Dating / relationships",
             "Men should improve themselves",
             "Equality / partnership", "Commentary / opinion",
             "4%", "62%", "0%", "11%", "85%"],
        ],
        col_widths=[2.6, 2.6, 3.0, 2.6, 2.4, 1.2, 1.2, 1.2, 1.2, 1.2],
        first_col_bold=True,
    )

    _add_heading(doc, "13.  Four entry-point archetypes", level=3)
    _add_para(doc,
        "Reading down the per-creator profiles, four clusters appear in "
        "this sample. Each archetype is a recurring topic + narrative + "
        "rhetoric + sexism combination, and is one of the ways an "
        "audience lands in the masculinity ecosystem."
    )

    _add_callout(doc,
        "Archetype 1 — Grievance and dominance",
        [
            "Creators: Andrew Kibe, Eric Amunga (Amerix), Wizarab.",
            "Narrative: \"men are disadvantaged victims\" / \"men should "
            "dominate, lead\". Topics: gender debate, dating / "
            "relationships. Frames: traditional gender order, male "
            "victimhood, female blame. Rhetoric: advice / instruction "
            "and commentary / opinion. Hostile-sexism rate: 41–74%.",
            "Audience signature: high opposes share (Amerix opposes "
            "55%); uptake splits between intensifies and challenges. "
            "This is the manosphere-aligned entry point.",
        ],
    )
    _add_callout(doc,
        "Archetype 2 — Provider and traditional order",
        [
            "Creators: Shola, Agba John Doe.",
            "Narrative: \"men should provide and succeed\" / \"men "
            "should protect women and children\". Topics: dating / "
            "relationships, money / status, marriage / family. "
            "Frames: traditional gender order, provider pressure, "
            "female blame. Rhetoric: commentary / opinion. "
            "Hostile-sexism rate: ~46%.",
            "Audience signature: high supports rate and high intensifies "
            "uptake. This is the soft-regressive entry point — gender "
            "hierarchy framed through wealth and bridewealth.",
        ],
    )
    _add_callout(doc,
        "Archetype 3 — Healing and emotional growth",
        [
            "Creators: Onyango Otieno (Rixpoet), Eddy Kimani, Banky "
            "Wellington.",
            "Narrative: \"men should be emotionally open\" / \"men "
            "should improve themselves\". Topics: mental health, "
            "fatherhood / parenting, gender debate. Frames: trauma "
            "healing, self-improvement, equality / partnership. "
            "Rhetoric: testimony, personal story, motivational speech. "
            "Hostile-sexism rate: ~0%.",
            "Audience signature: high supports and extends-original-"
            "message uptake; audience reports personal disclosure and "
            "learning / reflection at rates not observed under the other "
            "archetypes in this sample. The healing register is the "
            "clearest counter-narrative signature in the data.",
        ],
    )
    _add_callout(doc,
        "Archetype 4 — Partnership and fatherhood",
        [
            "Creators: Banky Wellington, Deyemi Okanlawon, Ebuka "
            "Obi-Uchendu, Philip Karanja.",
            "Narrative: \"men should be equal partners\" / \"men should "
            "protect women and children\". Topics: marriage / family, "
            "fatherhood / parenting. Frames: equality / partnership, "
            "fatherhood responsibility. Rhetoric: commentary / opinion "
            "and advice / instruction. Hostile-sexism rate: ~3%.",
            "Audience signature: Banky's audience supports and extends; "
            "Deyemi's audience opposes at 85%. Partnership and "
            "fatherhood content can be received as moralism unless it "
            "is paired with the healing register (Archetype 3).",
        ],
    )

    # ====================================================================
    # PART D · AUDIENCE FINDINGS
    # ====================================================================
    _add_heading(doc, "Part D · Audience reception (Scope Component 3)",
                 level=2)

    _add_heading(doc, "14.  Audience stance", level=3)
    _add_para(doc,
        "Stance is coded toward the original post, not toward gender "
        "in general. Per Section 12: Rixpoet's audience supports at "
        "88%; Amerix attracts the most opposition in Kenya. In Nigeria, "
        "Deyemi's audience opposes at 85% (68 of 80 comments)."
    )

    _add_heading(doc, "15.  Audience uptake", level=3)
    _add_para(doc,
        "Each comment is coded for how it relates to the post: "
        "intensifies the message (pushes it further), extends it "
        "(agrees and adds reasoning), softens it, or challenges it. "
        "This separates passive agreement from active amplification."
    )
    _add_table(doc,
        header=["Creator", "Country", "Intensifies", "Extends",
                "Challenges", "Comments (n)"],
        rows=[
            ["Amerix",  "Kenya",   "53 (38%)", "4 (3%)",   "76 (54%)", "140"],
            ["Eddy",    "Kenya",   "49 (45%)", "42 (38%)", "13 (12%)", "110"],
            ["Rixpoet", "Kenya",   "32 (35%)", "46 (50%)", "6 (7%)",   "92"],
            ["Kibe",    "Kenya",   "29 (41%)", "4 (6%)",   "26 (37%)", "70"],
            ["Agba JD", "Nigeria", "68 (42%)", "8 (5%)",   "61 (38%)", "161"],
            ["Banky W.","Nigeria", "39 (35%)", "41 (37%)", "12 (11%)", "110"],
            ["Deyemi",  "Nigeria", "37 (46%)", "2 (2%)",   "38 (48%)", "80"],
            ["Shola",   "Nigeria", "37 (56%)", "4 (6%)",   "24 (36%)", "66"],
        ],
        col_widths=[2.8, 2.0, 2.6, 2.6, 2.6, 2.4],
        first_col_bold=True,
    )

    _add_heading(doc, "16.  Perceived impact", level=3)
    _add_para(doc,
        "Each comment is coded for the impact it reports: validation, "
        "learning / reflection, personal disclosure, advice to others, "
        "resistance / pushback, emotional support, entertainment / "
        "humour, or no clear impact."
    )
    _add_bullets(doc, [
        "Kenya — Rixpoet drives personal disclosure and learning / "
        "reflection. Amerix drives resistance and pushback.",
        "Nigeria — Banky Wellington drives learning / reflection and "
        "emotional support. Deyemi drives resistance and pushback.",
    ])
    _add_callout(doc,
        "Key contrast",
        [
            "Kenya: validation (155) > resistance / pushback (100) > "
            "learning / reflection (69). Most audience engagement here "
            "is acknowledgement or amplification.",
            "Nigeria: resistance / pushback (150) > validation (113) > "
            "learning / reflection (78). Nigerian threads in this "
            "sample appear closer to organised counter-publics than to "
            "uniform echo chambers — a sample-level observation, but "
            "one that matters for playbook framing.",
            "Personal disclosure is a Kenya-specific signature in the "
            "sample (15 vs 3 comments) and appears almost entirely "
            "under the healing register, where it co-occurs with "
            "learning / reflection and emotional support.",
        ],
    )

    _add_heading(doc, "17.  Emotion landscape", level=3)
    _add_para(doc,
        "Anger and contempt cluster on regressive creators. Hope, "
        "admiration, and concern cluster on progressive creators. "
        "Sadness is rare but meaningful and appears almost exclusively "
        "under healing-register content (Rixpoet, Banky Wellington)."
    )

    _add_heading(doc, "18.  Audience argument types", level=3)
    _add_para(doc,
        "Kenya's audience leans on personal experience (120 of 412 = "
        "29%), driven by Rixpoet's mental-health register. Nigeria's "
        "audience leans on generalisation (167 of 417 = 40%), "
        "consistent with the dating / relationships topic mix."
    )

    _add_heading(doc, "19.  Audience-side sexism", level=3)
    _add_para(doc,
        "Audiences both reproduce and challenge sexism in the original "
        "posts. Hostile sexism in audience comments concentrates under "
        "the regressive-dominant creators who already produce it: "
        "Amerix's audience reproduces it; Rixpoet's audience codes 0%. "
        "In Nigeria the same pattern holds under Agba John Doe and "
        "Shola, with Banky's audience near zero."
    )

    # ====================================================================
    # PART E · LINKED CONTENT–AUDIENCE
    # ====================================================================
    _add_heading(doc, "Part E · Linked content–audience insights",
                 level=2)

    _add_heading(doc, "20.  Frame gap — where audiences open or close "
                       "frames", level=3)
    _add_para(doc,
        "The frame comparison in Section 5 shows where audiences are "
        "doing more than reflecting creators. Three patterns are large "
        "enough to flag for the playbook."
    )
    _add_bullets(doc, [
        "Kenya — content vs audience track each other on traditional "
        "gender order (79 → 58: audience deflates) and trauma healing "
        "(59 → 62: roughly reciprocal under Rixpoet and Eddy). The "
        "audience opens self-improvement and provider pressure more "
        "often than the content does (see §5).",
        "Nigeria — audience inflates traditional gender order "
        "(61 → 111) and female blame (75 → 82). The audience does the "
        "regressive frame-keeping work even where the content does not "
        "foreground these frames most.",
        "Nigeria — violence / safety appears in 39 audience comments "
        "under only 3 content snippets (almost all Deyemi Okanlawon). "
        "This is an unmet demand for GBV / safety conversation: the "
        "audience is asking for it; creators in this sample are not "
        "leading it.",
        "Trauma healing is concentrated and reciprocal in Kenya "
        "(content 59 / audience 62, almost entirely Rixpoet and Eddy). "
        "It does not exist in the Nigerian set (content 5 / audience 1). "
        "This is a counter-narrative gap rather than saturation.",
    ])

    _add_heading(doc, "21.  Where audiences amplify regressive frames",
                 level=3)
    _add_para(doc,
        "By creator, the gap between intensifies-the-message and "
        "challenges-the-message uptake (Section 15) shows where the "
        "audience is doing amplification work and where it is doing "
        "resistance work. Amerix and Deyemi sit on opposite sides of "
        "the same dynamic: high challenge rates against very different "
        "content. Banky and Rixpoet show the inverse — high extends "
        "shares, low challenge."
    )

    # ====================================================================
    # PART F · TOWARD THE PLAYBOOK
    # ====================================================================
    _add_heading(doc, "Part F · Toward the playbook (Component 4 bridge)",
                 level=2)

    _add_heading(doc, "22.  Counter-narrative levers visible in the data",
                 level=3)
    _add_para(doc,
        "Four narrative-and-rhetoric bundles already produce healthier "
        "audience signatures in this sample. Each is a candidate the "
        "playbook can lean on."
    )
    _add_callout(doc,
        "Lever 1 — Healing / testimony register",
        [
            "Pattern: personal story plus testimony rhetoric, trauma-"
            "healing frame, \"men should be emotionally open\" narrative.",
            "Why it works in the data: produces the highest rates of "
            "audience personal disclosure and learning / reflection. "
            "Engages male grievance without arguing against it.",
            "Where it works now: Rixpoet (KE), Eddy Kimani (KE), Banky "
            "Wellington (NG).",
            "Risk: thin Nigerian footprint — only Banky carries the "
            "register. The playbook should look for additional Nigerian "
            "creators in this register.",
        ],
    )
    _add_callout(doc,
        "Lever 2 — Fatherhood-led equality",
        [
            "Pattern: commentary / opinion plus advice rhetoric, "
            "fatherhood-responsibility frame, \"men should be equal "
            "partners\" narrative.",
            "Why it works: positions equality as responsible-father "
            "identity rather than concession. Audiences tend to extend "
            "rather than intensify.",
            "Where it works now: Banky Wellington, Philip Karanja. "
            "Deyemi Okanlawon uses this register but draws organised "
            "opposition (audience opposes 85%) — the content needs to "
            "be paired with the healing register to land.",
        ],
    )
    _add_callout(doc,
        "Lever 3 — Reframe the provider story without abandoning it",
        [
            "The most common Nigerian regressive narrative is \"men "
            "should provide and succeed\". The provider-pressure frame "
            "is real for adolescent boys and young men; direct attack "
            "is unlikely to land. The opening in the data is to "
            "decouple provider success from women's submission. Banky "
            "Wellington's content already does this; the playbook can "
            "scale that move.",
        ],
    )
    _add_callout(doc,
        "Lever 4 — Meet the unmet demand for GBV / safety conversation",
        [
            "In Nigeria, violence / safety appears in 39 audience "
            "comments under only 3 content snippets. The audience is "
            "asking for this conversation and the sample's creators are "
            "not leading it. The playbook can occupy this gap directly "
            "via Male Changemakers and progressive creators.",
        ],
    )

    # ====================================================================
    # PART G · APPENDICES
    # ====================================================================
    _add_heading(doc, "Part G · Appendices", level=2)

    _add_heading(doc, "23.  Definitions and computation", level=3)
    _add_para(doc,
        "Each coded variable and how its figures are computed. "
        "Percentages use the denominator stated in the table; "
        "percentages on the content side use the country's content-"
        "snippet n (KE = 394, NG = 381) unless otherwise noted; "
        "percentages on the audience side use the country's "
        "audience-comment n (KE = 412, NG = 417)."
    )
    _add_table(doc,
        header=["Variable", "What it captures", "Values",
                "Denominator", "Coded on"],
        rows=[
            ["Topic",
             "The main subject the snippet or comment is about.",
             "Mental health · Dating-relationships · Marriage-family · "
             "Gender debate · Fatherhood-parenting · Money-status · "
             "Violence-safety · Other.",
             "Country content n or audience n.",
             "Content + Audience"],
            ["Narrative",
             "The masculinity model the snippet pushes; \"none\" if "
             "descriptive only.",
             "Men are disadvantaged victims · Men should provide and "
             "succeed · Men should protect women and children · Men "
             "should improve themselves · Men should be emotionally "
             "open · Men should be equal partners · Men should "
             "dominate / lead · Other / none.",
             "Country content n.",
             "Content only"],
            ["Frame",
             "The dominant packaging of meaning around the topic.",
             "Traditional gender order · Trauma healing · Self-"
             "improvement · Provider pressure · Female blame · "
             "Equality-partnership · Fatherhood-responsibility · "
             "Violence-safety · Other.",
             "Country content n or audience n separately.",
             "Content + Audience"],
            ["Problem / Solution",
             "What the snippet says is wrong and what (if anything) it "
             "prescribes.",
             "Open category list; \"no clear problem\" / \"no clear "
             "solution\" allowed.",
             "Country content n.",
             "Content only"],
            ["Sentiment (3-axis)",
             "Separate sentiment toward men, women, and traditional "
             "norms.",
             "Positive · Negative · Neutral · Unclear.",
             "Country content n per target row.",
             "Content only"],
            ["Rhetoric",
             "Dominant rhetorical mode.",
             "Commentary / opinion · Advice / instruction · Debate / "
             "argument · Personal story · Testimony · Motivational "
             "speech · Religious / moral teaching · Humour / satire · "
             "Warning / threat · News / facts.",
             "Country content n.",
             "Content only"],
            ["Argument type",
             "How the claim is supported.",
             "Generalisation · Personal experience · Advice / "
             "instruction · Data / statistic · Religious / moral "
             "appeal · Other.",
             "Country content n or audience n.",
             "Content + Audience"],
            ["Sexism (3-level)",
             "Hostile, dehumanising, or benevolent / traditional "
             "sexism present in the snippet or comment.",
             "Hostile · Dehumanising · Benevolent-traditional · None.",
             "Country content n or audience n.",
             "Content + Audience"],
            ["Orientation",
             "Regressive vs progressive position on the masculinity "
             "axis.",
             "Regressive · Progressive · Mixed-unclear.",
             "Country content n.",
             "Content only (does NOT apply to audience)"],
            ["Stance (audience)",
             "Stance of the comment toward its parent post.",
             "Supports · Opposes · Neutral · Unclear.",
             "Country audience n.",
             "Audience only"],
            ["Uptake mode (audience)",
             "How the comment relates to the post.",
             "Intensifies · Extends · Softens · Challenges.",
             "Country audience n per creator (table 15).",
             "Audience only"],
            ["Perceived impact (audience)",
             "What kind of impact the commenter reports.",
             "Validation · Learning / reflection · Personal disclosure "
             "· Advice to others · Resistance / pushback · Emotional "
             "support · Entertainment / humour · No clear impact.",
             "Country audience n.",
             "Audience only"],
            ["Confidence",
             "LLM self-reported certainty on the row.",
             "High · Medium · Low.",
             "Country content n or audience n.",
             "All rows"],
        ],
        col_widths=[3.0, 3.6, 5.8, 2.8, 2.0],
        first_col_bold=True,
    )

    _add_para(doc,
        "Computation rules. All percentages are share = numerator ÷ "
        "denominator stated in the table. Where the denominator changes "
        "(per creator, per country, per target), the table reports the "
        "n in its own column. Counts and percentages are not weighted; "
        "they describe the curated sample only.",
        size=10, italic=True, color=COLOR_MUTED,
    )

    _add_heading(doc, "24.  Methodology", level=3)
    _add_para(doc,
        "Coding pipeline. Every row was coded by gpt-4o-mini using "
        "OpenAI Structured Outputs with a strict JSON schema that "
        "enumerates every allowed value. Temperature is fixed at 0. "
        "Responses are cached on disk by SHA-256 of (model, schema, "
        "system prompt, user prompt), so reruns are free and "
        "deterministic. Per-row uncertainty is captured by a "
        "confidence column (high / medium / low) and a verbatim "
        "evidence_quote field. The sentiment enum is four-valued "
        "(positive · negative · neutral · unclear). The LLM is the "
        "final coder."
    )
    _add_para(doc, "Coding completeness", italic=True, size=10)
    _add_table(doc,
        header=["Confidence", "KE audience", "KE content",
                "NG audience", "NG content"],
        rows=[
            ["High",   "321 (78%)", "365 (93%)", "335 (80%)", "340 (89%)"],
            ["Medium", "88 (21%)",  "27 (7%)",   "82 (20%)",  "41 (11%)"],
            ["Low",    "3",         "2",         "0",         "0"],
            ["Errors / fallbacks", "0", "0", "0", "0"],
        ],
        col_widths=[3.5, 3.0, 3.0, 3.0, 3.0],
        first_col_bold=True,
    )
    _add_para(doc,
        "Engagement metrics (view counts, impressions, comments, "
        "re-shares). Per Scope, engagement metrics are correlated with "
        "content variables in the final deck rather than in this "
        "exploratory document. View counts are not exposed on all "
        "platforms in the sample (e.g., X / Twitter does not surface "
        "view counts on older posts in all cases), so the deck reports "
        "metric-level analyses wherever the metric is available and "
        "flags the gap otherwise."
    )

    _add_heading(doc, "25.  Limitations, sampling, and how to read the "
                       "figures", level=3)
    _add_para(doc,
        "Sampling is purposive, not random. Eleven creators were "
        "selected to span the regressive ↔ progressive spectrum so that "
        "both ends of the masculinity discourse are observable in the "
        "data. The sample is therefore high-signal at both poles and "
        "is not a population estimate of \"Kenyan men\" or \"Nigerian "
        "men\". Corpus-level shares (the 43% / 35% / 31% headline "
        "numbers) describe what is in the sample; per-creator and "
        "per-archetype figures (Sections 2, 12, 13) are the primary "
        "unit of analysis."
    )
    _add_para(doc,
        "Three audiences are absent or partial. Ebuka Obi-Uchendu, "
        "Philip Karanja, and Wizarab have zero audience comments in the "
        "current sample (the scraped threads did not surface engaged "
        "audience text within the chosen window). Per-creator audience "
        "tables omit them; their content-side coding is included. "
        "Reading audience patterns at the country level should account "
        "for this."
    )
    _add_para(doc,
        "Language, sarcasm, and code-switching. Sarcasm, slang, and "
        "code-switching across English, Sheng, Swahili, and Pidgin "
        "carry a higher false-classification risk and are flagged with "
        "confidence = low (Section 24)."
    )
    _add_para(doc,
        "Single-coder limitation and validation plan. The LLM is the "
        "final coder in this exploratory pass. A human-coded recheck "
        "on a stratified random sample (target: 10% of rows per "
        "country, balanced across high / medium / low confidence) is "
        "scheduled in the focused pass and will produce a per-variable "
        "agreement rate against the LLM codes. Inter-rater reliability "
        "between human reviewers (Cohen's κ on a shared subset) will "
        "be reported alongside. The exploratory figures in this "
        "document should be read as conditional on the LLM coding "
        "until the validation pass is complete."
    )

    _add_heading(doc, "26.  Representative quotes", level=3)
    _add_para(doc, "Content side — Kenya", italic=True, size=10)
    _add_table(doc,
        header=["Orientation · Frame", "Creator", "Quote"],
        rows=[
            ["Regressive · social accountability", "Eric Amunga (Amerix)",
             "\"This woman had been assaulted by her husband after an "
             "evening quarrel because of food.\""],
            ["Regressive · female blame", "Eric Amunga (Amerix)",
             "\"Women are performers, and you, the man, are the audience.\""],
            ["Progressive · trauma healing", "Onyango Otieno (Rixpoet)",
             "\"the world out there people are struggling the same way "
             "we are struggling.\""],
            ["Progressive · male victimhood", "Onyango Otieno (Rixpoet)",
             "\"these are symptoms of trauma that you are exhibiting in "
             "your behavior\""],
            ["Regressive · self-improvement", "Eddy Kimani",
             "\"Because I'm the man. That was my thinking at the time.\""],
            ["Regressive · male victimhood", "Andrew Kibe",
             "\"your biggest bully was your mother.\""],
        ],
        col_widths=[4.2, 3.5, 8.5],
        first_col_bold=True,
    )

    _add_para(doc, "Content side — Nigeria", italic=True, size=10)
    _add_table(doc,
        header=["Orientation · Frame", "Creator", "Quote"],
        rows=[
            ["Regressive · trauma healing", "Banky Wellington",
             "\"I would never allow a woman to put me in that position "
             "again.\""],
            ["Progressive · fatherhood responsibility", "Banky Wellington",
             "\"fathers are supposed to give love and discipline, right?\""],
            ["Mixed · mixed / unclear", "Ebuka Obi-Uchendu",
             "\"the general perception, and I use general very strongly, "
             "is that Nigerian men are not that way.\""],
            ["Progressive · traditional gender order",
             "Ebuka Obi-Uchendu",
             "\"45% of men today aged 18 to 26 report that they've never "
             "approached a woman in person for dating.\""],
            ["Regressive · female blame", "Wizarab",
             "\"You are hoping he leaves his girlfriend for you one day. "
             "Keep dreaming\""],
            ["Progressive · equality / partnership", "Wizarab",
             "\"Tell her you're shy and smile.\""],
        ],
        col_widths=[4.2, 3.5, 8.5],
        first_col_bold=True,
    )

    _add_para(doc, "Audience side — Kenya", italic=True, size=10)
    _add_table(doc,
        header=["Stance · Frame", "Creator", "Quote"],
        rows=[
            ["Opposes · female blame", "Eddy Kimani",
             "\"this empowerment of women has challenged the men, and "
             "men have nothing left to challenge women with\""],
            ["Supports · trauma healing", "Onyango Otieno (Rixpoet)",
             "\"I try everyday to think positive about myself and after "
             "I get some money am planning to at least try counseling.\""],
            ["Supports · trauma healing", "Onyango Otieno (Rixpoet)",
             "\"I have been on the healing path for 15 years now and am "
             "glad am free!!\""],
            ["Supports · self-improvement", "Eddy Kimani",
             "\"peace is priceless to me now.\""],
            ["Supports · self-improvement", "Eric Amunga (Amerix)",
             "\"A man's desire for greatness is so strong that he will "
             "attempt to experience it through someone else.\""],
            ["Supports · provider pressure", "Andrew Kibe",
             "\"If you love cars, you don't park one in the garage for "
             "life and act like the showroom doesn't exist.\""],
        ],
        col_widths=[4.2, 3.5, 8.5],
        first_col_bold=True,
    )

    _add_para(doc, "Audience side — Nigeria", italic=True, size=10)
    _add_table(doc,
        header=["Stance · Frame", "Creator", "Quote"],
        rows=[
            ["Supports · self-improvement", "Banky Wellington",
             "\"the shame didn't just come from outside. They were "
             "already ashamed inside.\""],
            ["Supports · traditional gender order", "Banky Wellington",
             "\"Loving Banky's push back on many aspects of this "
             "conversation that was sliding into lazy generalising "
             "mindsets.\""],
            ["Opposes · violence / safety", "Deyemi Okanlawon",
             "\"You can't talk about rapists without talking about the "
             "accusers.\""],
            ["Opposes · male victimhood", "Deyemi Okanlawon",
             "\"you hopeless, horny degenerates calling yourselves men "
             "are the reason innocent boys have no say\""],
            ["Opposes · female blame", "Shola",
             "\"They are never, ever satisfied even when given "
             "everything\""],
            ["Opposes · female blame", "Shola",
             "\"women are then perceived to age faster, they often "
             "become desperate for marriage\""],
        ],
        col_widths=[4.2, 3.5, 8.5],
        first_col_bold=True,
    )

    _add_para(doc,
        "Source: coded CSVs from the Phase 1 LLM coding pipeline "
        "(gpt-4o-mini, schema v2, sentiment four-valued, temperature 0, "
        "deterministic cache). Section numbers map to the Gates "
        "Masculinity Project Scope of Work, Components 2 and 3.",
        size=9.5, color=COLOR_MUTED, italic=True,
    )

    doc.save(OUT_PATH)
    return doc


if __name__ == "__main__":
    build()
    print(f"wrote {OUT_PATH}")
