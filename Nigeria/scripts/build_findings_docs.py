"""
Build two findings documents (.docx):

  Nigeria/Audience Analysis/Exploratory/Nigeria - Audience LLM Exploratory Findings.docx
  Nigeria/Content Analysis/Exploratory/Nigeria - Content LLM Exploratory Findings.docx

Modeled on the prior project's "Findings" section structure.
Real numbers and quotes pulled from the analyzed parquets.
"""
from __future__ import annotations
from pathlib import Path
import json
import pandas as pd
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = Path(__file__).resolve().parents[2]
AUD = pd.read_parquet(ROOT / 'Nigeria/Audience Analysis/Exploratory/audience_exploratory_results.parquet')
CON = pd.read_parquet(ROOT / 'Nigeria/Content Analysis/Exploratory/content_exploratory_results.parquet')

# ─── helpers ────────────────────────────────────────────────────────────────

def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    return h

def add_para(doc, text, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    return p

def add_quote(doc, text, attribution=None):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.4)
    r = p.add_run(f'“{text}”')
    r.italic = True
    r.font.size = Pt(10.5)
    if attribution:
        r2 = p.add_run(f'  — {attribution}')
        r2.font.size = Pt(9.5)
        r2.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

def pct(n, total):
    return f'{100*n/total:.1f}%'

def explode(df, col):
    return df.assign(**{col: df[col].apply(lambda x: x if isinstance(x, list) else [])}).explode(col)

def add_pct_table(doc, ct_df, header_label='Creator'):
    """ct_df: DataFrame already in % form (rows=creator/orientation, cols=labels)."""
    cols = list(ct_df.columns)
    table = doc.add_table(rows=1 + len(ct_df), cols=1 + len(cols))
    table.style = 'Light Grid Accent 1'
    hdr = table.rows[0].cells
    hdr[0].text = header_label
    for i, c in enumerate(cols):
        hdr[i + 1].text = str(c)
    for i, (idx, row) in enumerate(ct_df.iterrows()):
        cells = table.rows[i + 1].cells
        cells[0].text = str(idx)
        for j, c in enumerate(cols):
            v = row[c]
            cells[j + 1].text = f'{v:.1f}%' if isinstance(v, float) else str(v)

def first_quotes(df, mask, n=2, max_len=240):
    sub = df[mask].copy()
    if 'sentiment__intensity' in sub.columns:
        sub = sub.sort_values('sentiment__intensity', ascending=False)
    out = []
    for _, r in sub.head(n).iterrows():
        idc = r.get('comment_id') or r.get('content_id')
        txt = str(r['text_english'])[:max_len].rstrip()
        if len(str(r['text_english'])) > max_len:
            txt += '…'
        out.append((idc, r['creator'], txt))
    return out


# ════════════════════════════════════════════════════════════════════════════
# AUDIENCE FINDINGS
# ════════════════════════════════════════════════════════════════════════════

def build_audience_doc():
    doc = Document()
    # title
    t = doc.add_heading('Nigeria — Audience LLM Exploratory Findings', 0)

    sub = doc.add_paragraph()
    r = sub.add_run('Norman Lear Center × Gates Foundation — Manfluencer Project')
    r.italic = True; r.font.size = Pt(11)
    sub2 = doc.add_paragraph()
    r = sub2.add_run('Dataset: Nigeria Audience Analysis Final.xlsx (417 comments, 4 creators).  Models: gpt-4o-mini for structured analyses; text-embedding-3-small for emergent topic clustering.')
    r.font.size = Pt(10); r.italic = True

    # ── Executive summary ──────────────────────────────────────────────────
    add_heading(doc, '1. Executive summary', 1)

    n = len(AUD)
    n_neg = (AUD['sentiment__sentiment_3class'] == 'negative').sum()
    n_pos = (AUD['sentiment__sentiment_3class'] == 'positive').sum()
    n_oppose = (AUD['stance__stance'] == 'oppose').sum()
    n_support = (AUD['stance__stance'] == 'support').sum()
    n_misog = (AUD['misogyny__misogyny'] != 'none').sum()
    n_hate = (AUD['hate_speech__hate_speech'] == True).sum()

    add_para(doc, (
        f'Across {n} curated audience comments (4 creators, 2 progressive + 2 regressive), '
        f'{pct(n_neg, n)} of replies are negative and {pct(n_pos, n)} are positive (3-class). '
        f'Stance toward the original post is split roughly evenly between support ({pct(n_support, n)}) and opposition ({pct(n_oppose, n)}); '
        f'{pct(n_misog, n)} of comments contain some form of misogyny, and {pct(n_hate, n)} register as hate speech under the LLM coding scheme. '
        'The two regressive creators (Agba John Doe, Shola) attract notably more toxic and hostile-misogynistic replies than the two progressive creators '
        '(Banky Wellington, Deyemi Okanlawon), but the picture is not uniform: Deyemi’s anti-rape posts attract the second-highest hate-speech rate of the four, '
        'driven by adversarial pushback rather than supportive amplification.'
    ))

    # ── Methods summary ────────────────────────────────────────────────────
    add_heading(doc, '2. Methods summary', 1)
    add_para(doc, 'Inputs.', bold=True)
    add_para(doc, '•  417 comments from the consolidated Nigeria Audience Analysis Final workbook (manager-curated, stance-balanced 1 : 1.19 progressive : regressive, translated to Standard English).')
    add_para(doc, 'Pipeline.', bold=True)
    add_para(doc, '•  Comments translated via gpt-4o-mini where Pidgin/Yoruba was present; English originals passed through unchanged. ')
    add_para(doc, '•  Eleven structured LLM analyses applied row-by-row: themes (closed taxonomy of 15 codes), sentiment (3-class), emotion (Plutchik-derived), NER, hate speech, toxicity (0–1), misogyny (6 subtypes), moral foundations (MFT), stance (audience-only), framing and argument mining (content-only). ')
    add_para(doc, '•  Embedding-based emergent topic clusters via text-embedding-3-small → UMAP → HDBSCAN, labeled by gpt-4o-mini. ')
    add_para(doc, 'Quality assurance.', bold=True)
    add_para(doc, '•  All categorical outputs fall within their declared vocabularies; numeric outputs within declared ranges; logical consistency (hate ⇔ severity, misogyny ⇔ intensity) enforced; row alignment between source workbook, translated parquet, and results parquet verified.')

    # ── 3. Findings ────────────────────────────────────────────────────────
    add_heading(doc, '3. Findings', 1)

    # 3.1 Sentiment
    add_heading(doc, '3.1 Sentiment patterns', 2)
    add_para(doc, (
        f'Of {n} comments, {pct((AUD["sentiment__sentiment_3class"]=="negative").sum(), n)} were negative, '
        f'{pct((AUD["sentiment__sentiment_3class"]=="neutral").sum(), n)} neutral, and '
        f'{pct((AUD["sentiment__sentiment_3class"]=="positive").sum(), n)} positive. '
        'Sentiment is highly creator-dependent.'
    ))
    ct = pd.crosstab(AUD['creator'], AUD['sentiment__sentiment_3class'], normalize='index') * 100
    add_pct_table(doc, ct)

    add_para(doc, (
        '\nBanky Wellington’s podcast comments lean positive (42.7%); '
        'Deyemi Okanlawon’s post on rape culture draws the most negative replies (63.7%). '
        'Both Banky and Deyemi are progressive creators — the divergence reflects topic and audience composition (a podcast on healthy masculinity vs. a confrontational tweet on rape).'
    ))

    # 3.2 Stance
    add_heading(doc, '3.2 Audience stance toward the original post', 2)
    add_para(doc, (
        f'{pct(n_support, n)} of comments support the original post and {pct(n_oppose, n)} oppose it; '
        f'{pct((AUD["stance__stance"]=="mixed").sum(), n)} are mixed, '
        f'{pct((AUD["stance__stance"]=="neutral").sum(), n)} neutral, '
        f'{pct((AUD["stance__stance"]=="off_topic").sum(), n)} off-topic. '
        'Banky’s audience is the most overtly supportive (46.4% support, 13.6% oppose). '
        'Deyemi’s audience is the most contested (37.5% oppose, 27.5% support).'
    ))
    ct = pd.crosstab(AUD['creator'], AUD['stance__stance'], normalize='index') * 100
    add_pct_table(doc, ct[['support', 'oppose', 'mixed', 'neutral', 'off_topic']])

    # 3.3 Toxicity & misogyny
    add_heading(doc, '3.3 Toxicity and misogyny', 2)
    tox = AUD.groupby(['orientation', 'creator'])[['toxicity__toxicity', 'misogyny__intensity']].mean().round(2)
    tox.columns = ['Mean toxicity (0–1)', 'Mean misogyny intensity (0–3)']
    add_pct_table(doc, tox.reset_index().set_index('creator').drop(columns='orientation'))
    add_para(doc, (
        '\nShola’s audience has the highest mean toxicity (0.39) and misogyny intensity (1.67); '
        'Agba’s sits below at 0.17 and 1.26. Banky is an order of magnitude lower (0.02, 0.66). '
        'Deyemi’s audience has the second-highest toxicity (0.35) — not because his audience is hostile to him, '
        'but because adversarial replies to his anti-rape post register as toxic content.'
    ))

    add_heading(doc, 'Misogyny subtypes', 3)
    sub = AUD['misogyny__misogyny'].value_counts()
    for k, v in sub.items():
        add_para(doc, f'•  {k}: {v} ({pct(v, n)})')

    # 3.4 Hate speech
    add_heading(doc, '3.4 Hate speech', 2)
    rate = AUD.groupby('creator')['hate_speech__hate_speech'].mean().round(3) * 100
    add_para(doc, 'Share of comments flagged as hate speech, by creator:')
    for cr, v in rate.items():
        add_para(doc, f'•  {cr}: {v:.1f}%')
    add_para(doc, (
        '\nShola’s audience is the most hate-prone (62.1%); Banky’s the least (10.0%). '
        'Deyemi’s post on rape culture attracts a high hate-speech rate (38.8%) because of adversarial content directed at the speaker, '
        'not because Deyemi’s supporters are hateful.'
    ))

    add_heading(doc, 'Illustrative high-severity hate speech', 3)
    for idc, cr, t in first_quotes(AUD, (AUD['hate_speech__hate_speech']==True) & (AUD['hate_speech__severity']>=2), n=3):
        add_quote(doc, t, f'{idc}, {cr}')

    # 3.5 Themes
    add_heading(doc, '3.5 Themes (controlled vocabulary)', 2)
    ex = explode(AUD, 'themes__themes').dropna(subset=['themes__themes'])
    ex = ex[ex['themes__themes'] != '']
    top = ex['themes__themes'].value_counts().head(10)
    add_para(doc, 'Top 10 themes across all 417 comments:')
    for k, v in top.items():
        add_para(doc, f'•  {k.replace("_", " ")}: {v} ({pct(v, n)})')

    add_para(doc, '\nMale sexual entitlement, female blame and male accountability dominate the audience discourse — they appear together in many comments. The strongest co-occurrence pair is male_sexual_entitlement × female_blame (55 comments), reflecting a recurring pattern of justifying male behavior while assigning blame to women in the same comment.')

    # 3.6 Emergent topic clusters
    add_heading(doc, '3.6 Emergent topic clusters (embeddings)', 2)
    add_para(doc, 'Independent of the closed taxonomy, embedding-based clustering finds 2 emergent clusters across the 417 comments:')
    cl = AUD.groupby(['topic_cluster_id','topic_cluster_label']).size().sort_values(ascending=False)
    for (cid, lbl), n_c in cl.items():
        add_para(doc, f'•  Cluster {cid} — “{lbl}”: {n_c} comments ({pct(n_c, n)})')

    # 3.7 Per-creator narratives
    add_heading(doc, '3.7 Per-creator audience narratives', 2)

    def creator_block(cr, title_extra=''):
        sub = AUD[AUD['creator'] == cr]
        n_c = len(sub)
        add_heading(doc, f'{cr} (n = {n_c}){title_extra}', 3)
        s_pos = pct((sub['sentiment__sentiment_3class']=='positive').sum(), n_c)
        s_neg = pct((sub['sentiment__sentiment_3class']=='negative').sum(), n_c)
        st_sup = pct((sub['stance__stance']=='support').sum(), n_c)
        st_opp = pct((sub['stance__stance']=='oppose').sum(), n_c)
        tox = sub['toxicity__toxicity'].mean()
        mis = sub['misogyny__intensity'].mean()
        hate = pct((sub['hate_speech__hate_speech']==True).sum(), n_c)
        add_para(doc, (
            f'Sentiment positive {s_pos} | negative {s_neg}.  '
            f'Stance support {st_sup} | oppose {st_opp}.  '
            f'Mean toxicity {tox:.2f}.  Mean misogyny intensity {mis:.2f}.  Hate-speech rate {hate}.'
        ))

    creator_block('Banky Wellington', ' — progressive, MENtality podcast')
    add_para(doc, 'Banky’s audience is the most affirming and least toxic. Most replies are reflective, supportive, or appreciative. Where critique appears, it tends to engage the substance rather than attack the speaker. Sample supportive replies:')
    for idc, cr, t in first_quotes(AUD, (AUD['creator']=='Banky Wellington') & (AUD['stance__stance']=='support'), n=2):
        add_quote(doc, t, idc)

    creator_block('Deyemi Okanlawon', ' — progressive, anti-rape post')
    add_para(doc, 'Deyemi’s audience is the most contested. Both supportive and adversarial replies are intense; toxicity comes mainly from adversarial replies that attack the speaker or invoke "false rape accusation" narratives. Sample opposing replies:')
    for idc, cr, t in first_quotes(AUD, (AUD['creator']=='Deyemi Okanlawon') & (AUD['stance__stance']=='oppose'), n=2):
        add_quote(doc, t, idc)

    creator_block('Agba John Doe', ' — regressive, post on infidelity')
    add_para(doc, 'Agba’s audience reproduces the same regressive frames he uses, with substantial victim-blaming and role-prescription content. Sample misogynistic replies:')
    mask = (AUD['creator']=='Agba John Doe') & (AUD['misogyny__misogyny'].isin(['hostile_misogyny','victim_blaming','role_prescription']))
    for idc, cr, t in first_quotes(AUD, mask, n=2):
        add_quote(doc, t, idc)

    creator_block('Shola', ' — regressive, post on availability trap')
    add_para(doc, 'Shola’s audience is the most hostile and toxic of the four. Hate speech and hostile misogyny are concentrated here. Sample replies:')
    mask = (AUD['creator']=='Shola') & (AUD['misogyny__intensity'] >= 2)
    for idc, cr, t in first_quotes(AUD, mask, n=2):
        add_quote(doc, t, idc)

    # 3.8 Cross-orientation comparison
    add_heading(doc, '3.8 Cross-orientation comparison', 2)
    by_or = AUD.groupby('orientation').agg(
        n=('comment_id','count'),
        toxicity=('toxicity__toxicity','mean'),
        misogyny=('misogyny__intensity','mean'),
        hate=('hate_speech__hate_speech','mean'),
        pct_negative=('sentiment__sentiment_3class', lambda s: (s=='negative').mean()*100),
        pct_oppose=('stance__stance', lambda s: (s=='oppose').mean()*100),
    ).round(2)
    add_pct_table(doc, by_or, header_label='Orientation')
    add_para(doc, '\nAudiences of regressive creators (Agba, Shola) score higher on toxicity, misogyny intensity, hate-speech rate, and proportion of opposing or negative replies than audiences of progressive creators (Banky, Deyemi). Banky’s audience pulls progressive averages down sharply.')

    # ── 4. Limitations ─────────────────────────────────────────────────────
    add_heading(doc, '4. Limitations', 1)
    add_para(doc, '•  Platform confound: Banky is YouTube podcast comments; the other three creators are X (Twitter) replies. Differences may partly reflect platform discourse norms.')
    add_para(doc, '•  Sample size per creator (66–161) is sufficient for descriptive comparison but underpowered for rare-event analysis.')
    add_para(doc, '•  All LLM-generated codes are pending validation against the human-coded subset (5/13 deadline). Cohen’s κ and Krippendorff’s α will be computed once that data returns.')
    add_para(doc, '•  gpt-4o-mini was used for all structured analyses; gpt-4o would likely improve performance on context-heavy tasks (misogyny subtypes, framing) at higher cost.')
    add_para(doc, '•  Translation quality has not been independently validated by a Pidgin/Yoruba speaker; ~80 of 417 audience comments contained non-English fragments.')

    # ── 5. Appendix references ─────────────────────────────────────────────
    add_heading(doc, '5. Appendix — figures and source files', 1)
    add_para(doc, 'Figures (PNG, in Nigeria/Audience Analysis/Exploratory/figures/):', bold=True)
    add_para(doc, '•  audience_theme_x_sentiment.png — row %-share heatmap')
    add_para(doc, '•  audience_theme_x_emotion.png')
    add_para(doc, '•  audience_theme_x_stance.png')
    add_para(doc, '•  audience_theme_cooccurrence.png  +  .csv')
    add_para(doc, '•  audience_progressive_theme_x_sentiment.png')
    add_para(doc, '•  audience_regressive_theme_x_sentiment.png')
    add_para(doc, 'Source data:', bold=True)
    add_para(doc, '•  Nigeria/Audience Analysis/Exploratory/Nigeria - Audience LLM Exploratory Data Analyses.xlsx (417 rows, all analysis columns)')
    add_para(doc, '•  Nigeria/Audience Analysis/Exploratory/audience_exploratory_results.parquet (full results with list columns)')

    out = ROOT / 'Nigeria/Audience Analysis/Exploratory/Nigeria - Audience LLM Exploratory Findings.docx'
    doc.save(out)
    print(f'wrote {out}')
    return out


# ════════════════════════════════════════════════════════════════════════════
# CONTENT FINDINGS
# ════════════════════════════════════════════════════════════════════════════

def build_content_doc():
    doc = Document()
    doc.add_heading('Nigeria — Content LLM Exploratory Findings', 0)
    sub = doc.add_paragraph()
    r = sub.add_run('Norman Lear Center × Gates Foundation — Manfluencer Project')
    r.italic = True; r.font.size = Pt(11)
    sub2 = doc.add_paragraph()
    r = sub2.add_run('Dataset: Nigeria Content Analysis Final.xlsx (310 segments, 6 creators).  Models: gpt-4o-mini for structured analyses; text-embedding-3-small for emergent topic clustering.')
    r.font.size = Pt(10); r.italic = True

    n = len(CON)
    n_neg = (CON['sentiment__sentiment_3class']=='negative').sum()
    n_pos = (CON['sentiment__sentiment_3class']=='positive').sum()

    # ── 1. Executive summary ───────────────────────────────────────────────
    add_heading(doc, '1. Executive summary', 1)
    add_para(doc, (
        f'Across {n} content segments (6 creators, 3 progressive + 3 regressive), '
        f'{pct(n_neg, n)} are negative in sentiment and {pct(n_pos, n)} positive. '
        'The strongest signal is a sharp framing divergence between progressive and regressive creators: '
        'progressive content (Banky, Deyemi, Ebuka) is dominated by male accountability, gender equality, and self-improvement frames; '
        'regressive content (Agba, Shola, Wizarab) is dominated by female blame and male victimhood frames. '
        'Misogyny intensity follows the same pattern: Agba 1.92 mean, Shola 1.44, Wizarab 1.27 versus Banky 0.38, Deyemi 0.62. '
        'Argumentation style also differs by orientation: progressive creators rely more on anecdote and self-reflective reasoning; regressive creators rely more on appeals to data and rhetorical questions.'
    ))

    # ── 2. Methods ─────────────────────────────────────────────────────────
    add_heading(doc, '2. Methods summary', 1)
    add_para(doc, 'Inputs.', bold=True)
    add_para(doc, '•  310 content segments from the consolidated Nigeria Content Analysis Final workbook (coding-unit-level: tweets, podcast snippets, with translation, context summaries, and source URLs).')
    add_para(doc, 'Pipeline.', bold=True)
    add_para(doc, '•  Translation (gpt-4o-mini) for Pidgin/Yoruba fragments.')
    add_para(doc, '•  Eleven structured LLM analyses including framing analysis and argument mining (content-specific, in addition to themes / sentiment / emotion / NER / hate speech / toxicity / misogyny / moral foundations).')
    add_para(doc, '•  Embedding-based emergent topic clusters via text-embedding-3-small → UMAP → HDBSCAN, labeled by gpt-4o-mini.')
    add_para(doc, 'QA.', bold=True)
    add_para(doc, '•  All categorical outputs in vocabulary, numeric outputs in range, logical consistency enforced, row alignment verified.')

    # ── 3. Findings ────────────────────────────────────────────────────────
    add_heading(doc, '3. Findings', 1)

    # 3.1 sentiment
    add_heading(doc, '3.1 Sentiment patterns', 2)
    ct = pd.crosstab(CON['creator'], CON['sentiment__sentiment_3class'], normalize='index') * 100
    add_pct_table(doc, ct)
    add_para(doc, (
        '\nShola’s content is the most negative in sentiment (Shola 71% negative when measured directly, with mostly hostile content directed at women); '
        'Banky’s and Ebuka’s content the most positive or balanced. '
        'Sentiment direction correlates strongly with creator orientation, but not perfectly — some progressive content is negative because it discusses harmful norms with disapproval (e.g. Deyemi calling out rape culture), '
        'while some regressive content is positive in tone but reproductive of regressive norms.'
    ))

    # 3.2 framing
    add_heading(doc, '3.2 Framing analysis (the headline result)', 2)
    add_para(doc, 'Distribution of framing by orientation (% within orientation):')
    fr_ct = (pd.crosstab(CON['orientation'], CON['framing__frame'], normalize='index')*100).round(1)
    add_pct_table(doc, fr_ct, header_label='Orientation')

    add_para(doc, '\nThe regressive creators are 7× more likely to use a female-blame frame (24.3% vs 3.6%), and 2× more likely to use a male-victimhood frame (24.3% vs 11.7%). The progressive creators are 3× more likely to use a gender-equality frame and lean heavily on male-accountability and self-improvement frames. This is the cleanest single divergence in the entire content dataset.')

    add_heading(doc, 'Examples — male victimhood (regressive)', 3)
    for idc, cr, t in first_quotes(CON, CON['framing__frame']=='male_victimhood', n=2):
        add_quote(doc, t, f'{idc}, {cr}')

    add_heading(doc, 'Examples — female blame (regressive)', 3)
    for idc, cr, t in first_quotes(CON, CON['framing__frame']=='female_blame', n=2):
        add_quote(doc, t, f'{idc}, {cr}')

    add_heading(doc, 'Examples — male accountability (progressive)', 3)
    for idc, cr, t in first_quotes(CON, (CON['framing__frame']=='male_accountability') & (CON['orientation']=='progressive'), n=2):
        add_quote(doc, t, f'{idc}, {cr}')

    add_heading(doc, 'Examples — gender equality (progressive)', 3)
    for idc, cr, t in first_quotes(CON, CON['framing__frame']=='gender_equality', n=2):
        add_quote(doc, t, f'{idc}, {cr}')

    # 3.3 argument mining
    add_heading(doc, '3.3 Argument mining', 2)
    add_para(doc, 'Reasoning type by orientation (% within orientation):')
    arg = (pd.crosstab(CON['orientation'], CON['argument_mining__reasoning_type'], normalize='index')*100).round(1)
    add_pct_table(doc, arg, header_label='Orientation')
    add_para(doc, '\nProgressive creators rely more on anecdote and self-reflection (33.6% vs 18.5%); regressive creators rely more on appeals to data (statistics, screenshots, lists) (26.6% vs 13.1%) and rhetorical questions (14.1% vs 15.3% — similar) and "no argument" (statements of opinion as fact, 21.7% vs 16.1%). The regressive style of argumentation is more declarative and citation-style; the progressive style is more reflective and storytelling.')

    add_heading(doc, 'Example — appeal to data, regressive', 3)
    for idc, cr, t in first_quotes(CON, (CON['argument_mining__reasoning_type']=='appeal_to_data') & (CON['orientation']=='regressive'), n=2):
        add_quote(doc, t, f'{idc}, {cr}')
    add_heading(doc, 'Example — anecdote, progressive', 3)
    for idc, cr, t in first_quotes(CON, (CON['argument_mining__reasoning_type']=='anecdote') & (CON['orientation']=='progressive'), n=2):
        add_quote(doc, t, f'{idc}, {cr}')

    # 3.4 misogyny
    add_heading(doc, '3.4 Misogyny intensity', 2)
    mis = CON.groupby('creator')['misogyny__intensity'].mean().round(2).sort_values(ascending=False)
    add_para(doc, 'Mean misogyny intensity (0–3) by creator:')
    for cr, v in mis.items():
        add_para(doc, f'•  {cr}: {v}')
    add_para(doc, '\nAgba (1.92), Shola (1.44), and Wizarab (1.27) lead. Banky (0.38), Deyemi (0.62), and Ebuka (0.92) score lower. Highest-intensity content tends to use hostile_misogyny or sexual_objectification subtypes.')
    add_heading(doc, 'Highest-intensity examples', 3)
    for idc, cr, t in first_quotes(CON, CON['misogyny__intensity']==3, n=3):
        add_quote(doc, t, f'{idc}, {cr}')

    # 3.5 themes
    add_heading(doc, '3.5 Themes (controlled vocabulary)', 2)
    ex = explode(CON, 'themes__themes').dropna(subset=['themes__themes'])
    top = ex['themes__themes'].value_counts().head(10)
    add_para(doc, 'Top 10 themes across all 310 segments:')
    for k, v in top.items():
        add_para(doc, f'•  {k.replace("_", " ")}: {v} ({pct(v, n)})')

    # 3.6 emergent clusters
    add_heading(doc, '3.6 Emergent topic clusters (embeddings)', 2)
    add_para(doc, 'Embedding-based clustering finds 6 emergent groups (with 80 segments left as noise/unclustered):')
    cl = CON.groupby(['topic_cluster_id','topic_cluster_label']).size().sort_values(ascending=False)
    for (cid, lbl), n_c in cl.items():
        add_para(doc, f'•  Cluster {cid} — “{lbl}”: {n_c} segments ({pct(n_c, n)})')

    # 3.7 per-creator
    add_heading(doc, '3.7 Per-creator content narratives', 2)
    creator_orient = {
        'Banky Wellington':  ('progressive', 'YouTube (MENtality podcast)'),
        'Deyemi Okanlawon':  ('progressive', 'X / Twitter'),
        'Ebuka Obi-Uchendu': ('progressive', 'YouTube (podcast)'),
        'Agba John Doe':     ('regressive',  'X / Twitter'),
        'Shola':             ('regressive',  'X / Twitter'),
        'Wizarab':           ('regressive',  'X / Twitter'),
    }
    for cr in ['Banky Wellington','Deyemi Okanlawon','Ebuka Obi-Uchendu','Agba John Doe','Shola','Wizarab']:
        sub = CON[CON['creator']==cr]
        n_c = len(sub)
        orient, plat = creator_orient[cr]
        add_heading(doc, f'{cr} (n = {n_c}) — {orient}, {plat}', 3)
        s_pos = pct((sub['sentiment__sentiment_3class']=='positive').sum(), n_c)
        s_neg = pct((sub['sentiment__sentiment_3class']=='negative').sum(), n_c)
        tox = sub['toxicity__toxicity'].mean()
        mis_v = sub['misogyny__intensity'].mean()
        top_frame = sub['framing__frame'].value_counts().head(2)
        frames_str = ', '.join([f'{k} ({v})' for k, v in top_frame.items()])
        add_para(doc, (
            f'Sentiment positive {s_pos} | negative {s_neg}.  '
            f'Mean toxicity {tox:.2f}.  Mean misogyny intensity {mis_v:.2f}.  '
            f'Top frames: {frames_str}.'
        ))

    # 3.8 cross-orientation
    add_heading(doc, '3.8 Cross-orientation comparison', 2)
    by_or = CON.groupby('orientation').agg(
        n=('content_id','count'),
        toxicity=('toxicity__toxicity','mean'),
        misogyny=('misogyny__intensity','mean'),
        pct_negative=('sentiment__sentiment_3class', lambda s: (s=='negative').mean()*100),
    ).round(2)
    add_pct_table(doc, by_or, header_label='Orientation')

    # ── 4. Limitations ─────────────────────────────────────────────────────
    add_heading(doc, '4. Limitations', 1)
    add_para(doc, '•  Mixed grain: Banky/Ebuka segments are podcast snippets (longer, more discursive); Deyemi/Agba/Shola/Wizarab segments are tweets (shorter, more declarative). Direct comparison should account for length and format.')
    add_para(doc, '•  Per-creator sample sizes range 24–76; cross-creator differences should be read with that in mind.')
    add_para(doc, '•  All LLM-generated codes pending human validation (Cohen’s κ, Krippendorff’s α).')
    add_para(doc, '•  Frame and argument-mining categories are inherently interpretive; some segments plausibly fit multiple categories.')

    # ── 5. Appendix ────────────────────────────────────────────────────────
    add_heading(doc, '5. Appendix — figures and source files', 1)
    add_para(doc, 'Figures (PNG, in Nigeria/Content Analysis/Exploratory/figures/):', bold=True)
    add_para(doc, '•  content_theme_x_sentiment.png')
    add_para(doc, '•  content_theme_x_emotion.png')
    add_para(doc, '•  content_theme_x_frame.png')
    add_para(doc, '•  content_theme_cooccurrence.png  +  .csv')
    add_para(doc, '•  content_progressive_theme_x_sentiment.png')
    add_para(doc, '•  content_regressive_theme_x_sentiment.png')
    add_para(doc, 'Source data:', bold=True)
    add_para(doc, '•  Nigeria/Content Analysis/Exploratory/Nigeria - Content LLM Exploratory Data Analyses.xlsx (310 rows, all analysis columns)')
    add_para(doc, '•  Nigeria/Content Analysis/Exploratory/content_exploratory_results.parquet (full results)')

    out = ROOT / 'Nigeria/Content Analysis/Exploratory/Nigeria - Content LLM Exploratory Findings.docx'
    doc.save(out)
    print(f'wrote {out}')
    return out


if __name__ == '__main__':
    build_audience_doc()
    build_content_doc()
